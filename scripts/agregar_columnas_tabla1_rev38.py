from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev37_(ZUJ)_22ago2026_Tabla1Reconstruida.docx")
OUTPUT = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev38_(ZUJ)_22ago2026_Tabla1ColumnasCompletas.docx")

METADATA = {
    "Hübner et al. (2024)": ("Alemania", "panadería; ML; pronóstico; desperdicio"),
    "Huber y Stuckenschmidt (2020)": ("Alemania", "retail; ML; calendario; demanda diaria"),
    "Abrar et al. (2024)": ("No especificado", "cadena de suministro; CNN; LSTM; XAI"),
    "Spiliotis et al. (2022)": ("Grecia", "SKU; demanda diaria; ML; métodos estadísticos"),
    "Makridakis et al. (2018)": ("Internacional", "M4; forecasting; ML; métodos híbridos"),
    "Hewamalage et al. (2021)": ("Australia", "RNN; series temporales; ARIMA; ETS"),
    "Giannopoulos et al. (2025)": ("Grecia", "demanda intermitente; ML; evaluación; Small Data"),
    "Borsato (2023)": ("Italia", "repostería; estética; personalización"),
    "Ab Karim et al. (2020)": ("Malasia", "pastelería; creatividad; experiencia"),
    "Ben Abdallah et al. (2022)": ("Túnez", "innovación; pastelería; competitividad"),
    "Mikalef et al. (2019)": ("Noruega", "big data; capacidades analíticas; innovación"),
    "Dubey et al. (2021)": ("Reino Unido / India", "IA; big data; desempeño operativo"),
    "Aprodu et al. (2025)": ("Rumania", "panadería; Industria 4.0; digitalización"),
    "Alekseeva et al. (2021)": ("Lituania", "adopción de IA; PYMES; barreras"),
    "Maldonado-Guzmán et al. (2018)": ("México", "innovación tecnológica; PYMES; desempeño"),
    "Cuevas-Vargas et al. (2016)": ("México", "TIC; PYMES; innovación"),
    "Maldonado-Guzmán y Garza-Reyes (2020)": ("México", "transformación digital; PYMES; datos"),
    "Saavedra-García y Tapia-Sánchez (2020)": ("México", "TIC; PYMES; competitividad"),
    "Romero-Hidalgo et al. (2021)": ("México", "gestión del conocimiento; innovación; datos"),
    "Poveda-Valverde y Fierro Barragán (2026)": ("Ecuador / América Latina", "IA; PYMES; adopción; barreras"),
}


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(6.6 if not bold else 6.9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def set_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(inches * 1440)))
    tc_width.set(qn("w:type"), "dxa")


doc = Document(SOURCE)
old = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "Autor(es) y año")
records = []
for row in old.rows[1:]:
    values = [cell.text.strip() for cell in row.cells]
    author = values[0]
    if author not in METADATA:
        raise ValueError(f"Metadata missing for {author}")
    country, keywords = METADATA[author]
    records.append((author, country, keywords, *values[1:]))

headers = ["Autor(es) y año", "País de origen", "Palabras clave", "Contexto y datos", "Metodología", "Hallazgo relevante", "Limitación o brecha", "Relación con la investigación"]
widths = [0.75, 0.6, 0.75, 0.76, 0.72, 0.98, 0.87, 0.98]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
table.autofit = False

for i, text in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], text, bold=True, fill="D9EAF7")
    set_width(table.rows[0].cells[i], widths[i])
for record in records:
    cells = table.add_row().cells
    for i, value in enumerate(record):
        set_cell_text(cells[i], value)
        set_width(cells[i], widths[i])

tr_pr = table.rows[0]._tr.get_or_add_trPr()
header = OxmlElement("w:tblHeader")
header.set(qn("w:val"), "true")
tr_pr.append(header)

old._element.addprevious(table._tbl)
old._element.getparent().remove(old._element)
doc.save(OUTPUT)
print(OUTPUT)
