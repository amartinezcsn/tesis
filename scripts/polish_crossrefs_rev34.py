from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC=Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev34_(ZUJ)_22ago2026.docx")
d=Document(DOC); body=d._body._element

phrases={
 'Tabla':' Los datos correspondientes se sintetizan en la Tabla ',
 'Figura':' La representación visual correspondiente se presenta en la Figura ',
 'Imagen':' La fuente visual correspondiente se identifica en la Imagen ',
 'Ecuacion':' La relación se formaliza en la ecuación ',
}

def lname(x): return x.tag.rsplit('}',1)[-1]

def ref_node(bookmark,result):
 f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),f' REF {bookmark} \\h ')
 r=OxmlElement('w:r'); t=OxmlElement('w:t'); t.text=str(result); r.append(t); f.append(r); return f

def find_ref_par(bookmark):
 for p in d.paragraphs:
  if any(bookmark in f.get(qn('w:instr'),'') for f in p._p.findall('.//'+qn('w:fldSimple'))): return p
 return None

def remove_generated(p,bookmark):
 if p is None: return
 for f in list(p._p.findall('./'+qn('w:fldSimple'))):
  if bookmark not in f.get(qn('w:instr'),''): continue
  prev=f.getprevious(); nxt=f.getnext()
  prevtxt=''.join(t.text or '' for t in prev.findall('.//'+qn('w:t'))) if prev is not None else ''
  nxttxt=''.join(t.text or '' for t in nxt.findall('.//'+qn('w:t'))) if nxt is not None else ''
  if prev is not None and prevtxt in phrases.values(): p._p.remove(prev)
  if nxt is not None and nxttxt=='.': p._p.remove(nxt)
  p._p.remove(f)

def append_ref(p,kind,bookmark,result):
 p.add_run(phrases[kind]); p._p.append(ref_node(bookmark,f'({result})' if kind=='Ecuacion' else result)); p.add_run('.')

def acceptable(p):
 st=p.style.name if p.style else ''; tx=p.text.strip()
 return len(tx)>=25 and st!='Caption' and not st.startswith('Heading') and not st.startswith('TOC') and not p._p.findall('.//'+qn('w:drawing'))

def next_context(element):
 children=list(body); pos=children.index(element)
 for node in children[pos+1:]:
  if lname(node)!='p': continue
  p=Paragraph(node,d._body)
  if acceptable(p) and not p.text.strip().startswith(('Tabla ','Figura ','Imagen ','Ecuación ')): return p
 return None

# Tables whose caption precedes the table are cited from the explanatory text that follows it.
for b in d._element.findall('.//'+qn('w:bookmarkStart')):
 name=b.get(qn('w:name'),'')
 if not name.startswith('xref_tabla_'): continue
 cap=b.getparent()
 while cap is not None and lname(cap)!='p': cap=cap.getparent()
 if cap is None: continue
 nxt=cap.getnext()
 while nxt is not None and lname(nxt)=='bookmarkEnd': nxt=nxt.getnext()
 if nxt is None or lname(nxt)!='tbl': continue
 dest=next_context(nxt)
 old=find_ref_par(name)
 if dest is None or (old is not None and old._p is dest._p): continue
 number=int(name.rsplit('_',1)[1])
 remove_generated(old,name); append_ref(dest,'Tabla',name,number)

# Equations 7-9 were inserted immediately after subsection labels; cite them from the following explanation.
for number in (7,8,9):
 name=f'xref_ecuacion_{number:02d}'
 target=next((b for b in d._element.findall('.//'+qn('w:bookmarkStart')) if b.get(qn('w:name'))==name),None)
 if target is None: continue
 table=target
 while table is not None and lname(table)!='tbl': table=table.getparent()
 dest=next_context(table) if table is not None else None
 old=find_ref_par(name)
 if dest is not None and (old is None or old._p is not dest._p):
  remove_generated(old,name); append_ref(dest,'Ecuacion',name,number)

# Rolling-Origin is cited in its own evaluation paragraph.
name='xref_ecuacion_10'; old=find_ref_par(name)
dest=next((p for p in d.paragraphs if p.text.strip().startswith('La evaluación se ejecutó mediante Rolling-Origin')),None)
if old is not None and dest is not None and old._p is not dest._p:
 remove_generated(old,name); append_ref(dest,'Ecuacion',name,10)

# Remove static equation mentions that duplicated the new dynamic REF fields.
replacements={
 'Matemáticamente, la función de pérdida asimétrica':('mediante la ecuación (1) del costo estimado:','mediante el costo esperado del modelo Newsvendor.'),
 'James et al.':('la ecuación () (EPE)','el error de predicción esperado (EPE)'),
 'Regresión Lasso':('representada matemáticamente en la ecuación (3):','representada matemáticamente mediante regularización L1.'),
 'Sin embargo, el modelo no interpreta':('usando la Ecuación 4:','mediante una proyección seno-coseno.'),
}
for start,(oldtxt,newtxt) in replacements.items():
 p=next((x for x in d.paragraphs if x.text.strip().startswith(start)),None)
 if p is None: continue
 # Preserve the generated REF by rebuilding it after the textual correction.
 number=next((n for n in range(1,5) if (find_ref_par(f'xref_ecuacion_{n:02d}') is not None and find_ref_par(f'xref_ecuacion_{n:02d}')._p is p._p)),None)
 if number is None: continue
 name=f'xref_ecuacion_{number:02d}'; remove_generated(p,name)
 corrected=p.text.replace(oldtxt,newtxt)
 p.clear(); p.add_run(corrected); append_ref(p,'Ecuacion',name,number)

# Combine the five metric references into one readable sentence.
metric_par=next((p for p in d.paragraphs if all(find_ref_par(f'xref_ecuacion_{n:02d}') is not None and find_ref_par(f'xref_ecuacion_{n:02d}')._p is p._p for n in range(11,16))),None)
if metric_par is not None:
 for n in range(11,16): remove_generated(metric_par,f'xref_ecuacion_{n:02d}')
 metric_par.add_run(' Las métricas y el criterio comparativo se formalizan en las ecuaciones ')
 for n in range(11,16):
  metric_par._p.append(ref_node(f'xref_ecuacion_{n:02d}',f'({n})'))
  metric_par.add_run(', ' if n<14 else (' y ' if n==14 else '.'))

# Combine adjacent references to the two result figures.
figpar=next((p for p in d.paragraphs if find_ref_par('xref_figura_08') is not None and find_ref_par('xref_figura_09') is not None and find_ref_par('xref_figura_08')._p is p._p and find_ref_par('xref_figura_09')._p is p._p),None)
if figpar is not None:
 remove_generated(figpar,'xref_figura_08'); remove_generated(figpar,'xref_figura_09')
 figpar.add_run(' La mejora frente a la línea base y la selección de modelos se muestran, respectivamente, en las Figuras ')
 figpar._p.append(ref_node('xref_figura_08','31')); figpar.add_run(' y ')
 figpar._p.append(ref_node('xref_figura_09','32')); figpar.add_run('.')

# Remove orphaned generic phrases left after moving a REF field between paragraphs.
for p in d.paragraphs:
 children=list(p._p)
 for idx,node in list(enumerate(children)):
  txt=''.join(t.text or '' for t in node.findall('.//'+qn('w:t')))
  if txt not in phrases.values(): continue
  nxt=children[idx+1] if idx+1<len(children) else None
  is_ref=nxt is not None and lname(nxt)=='fldSimple' and 'REF xref_' in nxt.get(qn('w:instr'),'')
  if is_ref: continue
  p._p.remove(node)
  if nxt is not None and ''.join(t.text or '' for t in nxt.findall('.//'+qn('w:t')))=='.' and nxt.getparent() is p._p:
   p._p.remove(nxt)

d.save(DOC)
print('crossrefs polished')
