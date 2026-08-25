from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC=Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev34_(ZUJ)_22ago2026.docx")
d=Document(DOC)

TABLE_PHRASE=' Los datos correspondientes se sintetizan en la Tabla '
EQ_PHRASE=' La relación se formaliza en la ecuación '

def field(bookmark,result):
 f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),f' REF {bookmark} \\h ')
 r=OxmlElement('w:r'); t=OxmlElement('w:t'); t.text=str(result); r.append(t); f.append(r); return f

def field_par(bookmark):
 for p in d.paragraphs:
  if any(bookmark in f.get(qn('w:instr'),'') for f in p._p.findall('./'+qn('w:fldSimple'))): return p

def remove_generated(p,bookmark,phrase):
 if p is None:return
 for f in list(p._p.findall('./'+qn('w:fldSimple'))):
  if bookmark not in f.get(qn('w:instr'),''):continue
  prev=f.getprevious(); nxt=f.getnext()
  if prev is not None and ''.join(t.text or '' for t in prev.findall('.//'+qn('w:t')))==phrase:p._p.remove(prev)
  if nxt is not None and ''.join(t.text or '' for t in nxt.findall('.//'+qn('w:t')))=='.' and nxt.getparent() is p._p:p._p.remove(nxt)
  p._p.remove(f)

def replace_text_with_ref(p,old,before_label,bookmark,result):
 text=p.text
 if old not in text:return
 left,right=text.split(old,1); p.clear(); p.add_run(left+before_label); p._p.append(field(bookmark,result)); p.add_run(right)

# Replace existing static table mentions with true cross-references.
jobs=[
 ('xref_tabla_01',1,'Con el propósito de sintetizar','Tabla 1','Tabla '),
 ('xref_tabla_05',5,'Para los fines de la investigación','Tabla 5','Tabla '),
 ('xref_tabla_06',6,'Con el propósito de documentar','Tabla 6','Tabla '),
 ('xref_tabla_11',11,'ARIMA fue la familia ganadora','la tabla siguiente','la Tabla '),
]
for bm,num,start,old,label in jobs:
 source=field_par(bm); dest=next((p for p in d.paragraphs if p.text.strip().startswith(start)),None)
 remove_generated(source,bm,TABLE_PHRASE)
 if dest is not None:replace_text_with_ref(dest,old,label,bm,num)

# Place the inflation-adjustment equation reference in the INPC explanation.
bm='xref_ecuacion_05'; source=field_par(bm)
dest=next((p for p in d.paragraphs if p.text.strip().startswith('Este nuevo conjunto de datos') and 'INPC' in p.text),None)
remove_generated(source,bm,EQ_PHRASE)
if dest is not None:
 dest.add_run(EQ_PHRASE); dest._p.append(field(bm,'(5)')); dest.add_run('.')

d.save(DOC)
print('targeted fixes complete')
