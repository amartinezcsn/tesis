from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC=Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev34_(ZUJ)_22ago2026.docx")
d=Document(DOC)
phrases={
 'Tabla':' Los datos correspondientes se sintetizan en la Tabla ',
 'Ecuacion':' La relación se formaliza en la ecuación ',
}

def ref(bookmark,result):
 f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),f' REF {bookmark} \\h ')
 r=OxmlElement('w:r'); t=OxmlElement('w:t'); t.text=str(result); r.append(t); f.append(r); return f

def direct_refs(p):
 return [f for f in p._p.findall('./'+qn('w:fldSimple')) if 'REF xref_' in f.get(qn('w:instr'),'')]

def find_par(bookmark):
 for p in d.paragraphs:
  if any(bookmark in f.get(qn('w:instr'),'') for f in direct_refs(p)): return p

def remove_generated(p,bookmark,phrase):
 if p is None: return
 for f in list(direct_refs(p)):
  if bookmark not in f.get(qn('w:instr'),''): continue
  prev=f.getprevious(); nxt=f.getnext()
  if prev is not None and ''.join(t.text or '' for t in prev.findall('.//'+qn('w:t')))==phrase: p._p.remove(prev)
  if nxt is not None and ''.join(t.text or '' for t in nxt.findall('.//'+qn('w:t')))=='.' and nxt.getparent() is p._p: p._p.remove(nxt)
  p._p.remove(f)

def append(p,phrase,bookmark,result):
 p.add_run(phrase); p._p.append(ref(bookmark,result)); p.add_run('.')

# Correct two remaining static equation mentions and preserve the dynamic fields.
for number,needle,replacement in [
 (2,'la ecuación () (EPE)','el error de predicción esperado (EPE)'),
 (4,'usando la Ecuación 4:','mediante una proyección seno-coseno.'),
]:
 p=next((p for p in d.paragraphs if needle in p.text),None)
 if p is not None:
  bm=f'xref_ecuacion_{number:02d}'; remove_generated(p,bm,phrases['Ecuacion'])
  text=p.text.replace(needle,replacement); p.clear(); p.add_run(text); append(p,phrases['Ecuacion'],bm,f'({number})')

# Consolidate the five metric references in Results.
metric=next((p for p in d.paragraphs if sum(1 for f in direct_refs(p) if any(f'xref_ecuacion_{n:02d}' in f.get(qn('w:instr'),'') for n in range(11,16)))==5),None)
if metric is not None:
 for n in range(11,16): remove_generated(metric,f'xref_ecuacion_{n:02d}',phrases['Ecuacion'])
 metric.add_run(' Las métricas y el criterio comparativo se formalizan en las ecuaciones ')
 for n in range(11,16):
  metric._p.append(ref(f'xref_ecuacion_{n:02d}',f'({n})'))
  metric.add_run(', ' if n<14 else (' y ' if n==14 else '.'))

# Relocate four table references from source notes or the next chapter to their explanatory paragraphs.
moves={
 'xref_tabla_01':('La Tabla 1 presenta',1),
 'xref_tabla_05':('La temperatura media mensual',5),
 'xref_tabla_06':('La integración de las fuentes',6),
 'xref_tabla_11':('ARIMA fue la familia ganadora',11),
}
for bm,(start,num) in moves.items():
 old=find_par(bm); dest=next((p for p in d.paragraphs if p.text.strip().startswith(start)),None)
 if old is not None and dest is not None and old._p is not dest._p:
  remove_generated(old,bm,phrases['Tabla']); append(dest,phrases['Tabla'],bm,num)

# Delete any generic phrase no longer followed by a REF field.
for p in d.paragraphs:
 children=list(p._p)
 for idx,node in list(enumerate(children)):
  txt=''.join(t.text or '' for t in node.findall('.//'+qn('w:t')))
  if txt not in phrases.values(): continue
  nxt=children[idx+1] if idx+1<len(children) else None
  if nxt is not None and nxt.tag==qn('w:fldSimple') and 'REF xref_' in nxt.get(qn('w:instr'),''): continue
  if node.getparent() is p._p: p._p.remove(node)
  if nxt is not None and ''.join(t.text or '' for t in nxt.findall('.//'+qn('w:t')))=='.' and nxt.getparent() is p._p: p._p.remove(nxt)

d.save(DOC)
print('final cross-reference cleanup complete')
