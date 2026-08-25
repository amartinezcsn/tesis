from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev34_(ZUJ)_22ago2026.docx")
doc = Document(DOC)
body = doc._body._element


def local(tag):
    return tag.rsplit('}', 1)[-1]


def next_bookmark_id():
    values = []
    for node in doc._element.findall('.//' + qn('w:bookmarkStart')):
        try:
            values.append(int(node.get(qn('w:id'))))
        except (TypeError, ValueError):
            pass
    return max(values or [0]) + 1


bookmark_id = next_bookmark_id()


def bookmark_nodes(parent, first, last, name):
    global bookmark_id
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    bookmark_id += 1
    parent.insert(parent.index(first), start)
    parent.insert(parent.index(last) + 1, end)


def run_element(text, source_rpr=None):
    run = OxmlElement('w:r')
    if source_rpr is not None:
        run.append(deepcopy(source_rpr))
    t = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    return run


def add_ref_field(paragraph, bookmark, cached_text):
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), f' REF {bookmark} \\h ')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = cached_text
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def paragraph_has_ref(paragraph, bookmark):
    for field in paragraph._p.findall('.//' + qn('w:fldSimple')):
        if bookmark in field.get(qn('w:instr'), ''):
            return True
    for instr in paragraph._p.findall('.//' + qn('w:instrText')):
        if bookmark in (instr.text or ''):
            return True
    return False


def append_reference(paragraph, kind, bookmark, number):
    if paragraph_has_ref(paragraph, bookmark):
        return
    phrases = {
        'Tabla': ' Los datos correspondientes se sintetizan en la Tabla ',
        'Figura': ' La representación visual correspondiente se presenta en la Figura ',
        'Imagen': ' La fuente visual correspondiente se identifica en la Imagen ',
        'Ecuacion': ' La relación se formaliza en la ecuación ',
    }
    paragraph.add_run(phrases[kind])
    cached = f'({number})' if kind == 'Ecuacion' else str(number)
    add_ref_field(paragraph, bookmark, cached)
    paragraph.add_run('.')


def nearest_context(element):
    children = list(body)
    pos = children.index(element)
    for node in reversed(children[:pos]):
        if local(node.tag) != 'p':
            continue
        p = Paragraph(node, doc._body)
        style = p.style.name if p.style is not None else ''
        text = p.text.strip()
        if not text or len(text) < 25:
            continue
        if style == 'Caption' or style.startswith('Heading') or style.startswith('TOC'):
            continue
        if node.findall('.//' + qn('w:drawing')) or node.findall('.//' + qn('w:pict')):
            continue
        if text.startswith(('Figura ', 'Tabla ', 'Imagen ', 'Ecuación ')):
            continue
        return p
    return None


def caption_target(paragraph, kind, ordinal):
    """Bookmark only the displayed caption number and return its cached value."""
    parent = paragraph._p
    for field in parent.findall('./' + qn('w:fldSimple')):
        instr = field.get(qn('w:instr'), '')
        if 'SEQ ' in instr:
            values = [t.text for t in field.findall('.//' + qn('w:t')) if t.text]
            number = values[-1] if values else str(ordinal)
            bookmark_nodes(parent, field, field, f'xref_{kind.lower()}_{ordinal:02d}')
            return number

    match = re.match(r'^(Figura|Tabla|Imagen)\s+(\d+)', paragraph.text.strip())
    if not match:
        return str(ordinal)
    number = match.group(2)
    prefix_len = match.start(2)
    run = next((r for r in parent.findall('./' + qn('w:r')) if ''.join(r.itertext()).strip()), None)
    if run is None:
        return number
    original = ''.join(t.text or '' for t in run.findall('.//' + qn('w:t')))
    idx = parent.index(run)
    rpr = run.find(qn('w:rPr'))
    before, after = original[:prefix_len], original[prefix_len + len(number):]
    parent.remove(run)
    before_run = run_element(before, rpr)
    number_run = run_element(number, rpr)
    after_run = run_element(after, rpr)
    parent.insert(idx, before_run)
    parent.insert(idx + 1, number_run)
    parent.insert(idx + 2, after_run)
    bookmark_nodes(parent, number_run, number_run, f'xref_{kind.lower()}_{ordinal:02d}')
    return number


# Caption targets and references (tables, figures, and source images).
counts = {'Tabla': 0, 'Figura': 0, 'Imagen': 0}
caption_refs = 0
for p in list(doc.paragraphs):
    if p.style.name != 'Caption':
        continue
    match = re.match(r'^(Tabla|Figura|Imagen)\b', p.text.strip())
    if not match:
        continue
    kind = match.group(1)
    counts[kind] += 1
    ordinal = counts[kind]
    bookmark = f'xref_{kind.lower()}_{ordinal:02d}'
    number = caption_target(p, kind, ordinal)
    context = nearest_context(p._p)
    if context is not None:
        append_reference(context, kind, bookmark, number)
        caption_refs += 1


# Equation targets: bookmark the complete marker in the third cell, e.g. (8).
equation_refs = 0
equation_number = 0
for table in doc.tables:
    if len(table.columns) != 3 or not table.rows:
        continue
    marker_p = table.cell(0, 2).paragraphs[0]
    xml = marker_p._p
    field_codes = [x.get(qn('w:instr'), '') for x in xml.findall('.//' + qn('w:fldSimple'))]
    field_codes += [''.join(x.itertext()) for x in xml.findall('.//' + qn('w:instrText'))]
    if not any('SEQ Ecuacion' in code for code in field_codes):
        continue
    equation_number += 1
    bookmark = f'xref_ecuacion_{equation_number:02d}'
    direct = [node for node in xml if local(node.tag) != 'pPr']
    if not direct:
        continue
    bookmark_nodes(xml, direct[0], direct[-1], bookmark)
    context = nearest_context(table._tbl)
    if context is not None:
        append_reference(context, 'Ecuacion', bookmark, equation_number)
        equation_refs += 1


# Fields were materialized with correct cached results; avoid repairs of legacy TOC hyperlinks.
update = doc.settings._element.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    doc.settings._element.append(update)
update.set(qn('w:val'), 'false')

doc.core_properties.comments = (
    'Rev34: referencias cruzadas REF vinculadas a marcadores para tablas, figuras, imágenes y ecuaciones.'
)
doc.save(DOC)
print(f'captions={sum(counts.values())}; caption_refs={caption_refs}; equations={equation_number}; equation_refs={equation_refs}')
