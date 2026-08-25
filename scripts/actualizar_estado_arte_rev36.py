from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches


SOURCE = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev35_(ZUJ)_22ago2026_PresupuestoAbastecimiento.docx")
OUTPUT = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev36_(ZUJ)_22ago2026_EstadoDelArteActualizado.docx")


def normalized(value):
    return " ".join(value.split())


def find_prefix(document, prefix):
    prefix = normalized(prefix)
    matches = [paragraph for paragraph in document.paragraphs if normalized(paragraph.text).startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph beginning with {prefix!r}; found {len(matches)}")
    return matches[0]


def find_exact(document, text):
    text = normalized(text)
    matches = [paragraph for paragraph in document.paragraphs if normalized(paragraph.text) == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph matching {text[:80]!r}; found {len(matches)}")
    return matches[0]


def replace_text(paragraph, text):
    props = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if props is not None:
        run._element.insert(0, props)


def insert_after(document, anchor, text, style=None):
    paragraph = document.add_paragraph(style=style or anchor.style)
    source_ppr = anchor._p.pPr
    if source_ppr is not None:
        paragraph._p.get_or_add_pPr().getparent().replace(paragraph._p.get_or_add_pPr(), deepcopy(source_ppr))
    paragraph.add_run(text)
    anchor._p.addnext(paragraph._p)
    return paragraph


def insert_reference_after(document, anchor, parts):
    paragraph = document.add_paragraph(style=anchor.style)
    source_ppr = anchor._p.pPr
    if source_ppr is not None:
        paragraph._p.get_or_add_pPr().getparent().replace(paragraph._p.get_or_add_pPr(), deepcopy(source_ppr))
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    for text, italic in parts:
        run = paragraph.add_run(text)
        run.italic = italic
    anchor._p.addnext(paragraph._p)
    return paragraph


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(SOURCE)

# Corrige una afirmación incompatible con la evidencia comparativa de la tesis.
replace_text(
    find_prefix(doc, "La literatura reciente destaca que los modelos predictivos basados"),
    "La literatura reciente muestra que los modelos de aprendizaje automático pueden capturar relaciones no lineales, estacionalidad y variables externas; sin embargo, su superioridad frente a los métodos estadísticos depende de las características de la serie, la disponibilidad de datos y el protocolo de evaluación (Makridakis et al., 2018; Spiliotis et al., 2022).",
)

# Nuevo bloque crítico sobre el problema técnico que sí corresponde al estudio.
anchor = find_prefix(doc, "La repostería creativa como sector económico especializado")
heading = anchor.insert_paragraph_before("Demanda intermitente, Small Data y comparación de modelos", style="Heading 3")
paragraph = insert_after(
    doc,
    heading,
    "Las series diarias de ventas y compras en negocios de pequeña escala suelen contener numerosos periodos sin movimiento y picos asociados a pedidos, fechas comerciales o reposiciones. Spiliotis et al. (2022) señalan que el pronóstico diario a nivel de SKU o de punto de venta enfrenta irregularidad e intermitencia, por lo que los modelos de aprendizaje automático deben compararse con métodos estadísticos establecidos y no asumirse superiores por su mayor complejidad.",
)
paragraph = insert_after(
    doc,
    paragraph,
    "La evidencia de las competencias M respalda este criterio de comparación. En M4, las combinaciones de métodos estadísticos y los enfoques híbridos obtuvieron resultados competitivos, mientras que los modelos puramente basados en aprendizaje automático no dominaron de forma generalizada (Makridakis et al., 2018). De manera consistente, Hewamalage et al. (2021) advierten que las redes neuronales recurrentes son alternativas competitivas, pero no soluciones universales, especialmente cuando las series individuales son cortas o heterogéneas.",
)
insert_after(
    doc,
    paragraph,
    "La revisión de Giannopoulos et al. (2025) confirma que la demanda intermitente requiere atención explícita al diseño de características, la partición temporal, el ajuste de hiperparámetros y las métricas de evaluación. Esta literatura sustenta que, en un contexto de Small Data, una validación Rolling-Origin y la comparación con líneas base reproducibles son más pertinentes que inferir superioridad a partir de un único modelo o una única partición.",
)

# Contexto latinoamericano actualizado y delimitado.
anchor = find_prefix(doc, "Investigaciones en México y América Latina")
latin_anchor = find_prefix(doc, "Finalmente, Romero-Hidalgo")
insert_after(
    doc,
    latin_anchor,
    "Como complemento regional, Poveda-Valverde y Fierro Barragán (2026), en una revisión sistemática de estudios sobre IA en PYMES latinoamericanas, identifican restricciones de infraestructura, presupuesto y talento especializado, y observan que las aplicaciones se concentran principalmente en logística, manufactura y telecomunicaciones. Este hallazgo refuerza la necesidad de estudios aplicados y reproducibles en microempresas de servicios alimentarios, donde la disponibilidad de datos es más limitada.",
)

# Síntesis y vacío alineados con el objetivo vigente.
replace_text(
    find_prefix(doc, "El análisis de la literatura evidencia que los modelos de aprendizaje automático"),
    "La literatura revisada muestra que el aprendizaje automático puede mejorar el pronóstico en determinados contextos de demanda y abastecimiento, pero también evidencia que los resultados dependen de la intermitencia de la serie, la granularidad, el volumen de datos, la inclusión de variables exógenas y el protocolo de validación.",
)
replace_text(
    find_prefix(doc, "Sin embargo, la mayoría de las investigaciones existentes"),
    "Los estudios de comparación más próximos se desarrollan con datos de SKU, minoristas o grandes colecciones de series, mientras que las revisiones sobre adopción de IA en PYMES se concentran en barreras organizacionales. Por ello, tales hallazgos no se trasladan automáticamente a una microempresa alimentaria que sólo dispone de registros internos de ventas y compras.",
)
replace_text(
    find_prefix(doc, "En contraste, existe una limitada evidencia empírica"),
    "La evidencia identificada es limitada respecto de comparaciones temporales reproducibles entre reglas empíricas, modelos estadísticos y aprendizaje automático para anticipar el importe diario de compras en microempresas de repostería con datos escasos e intermitentes.",
)
replace_text(
    find_prefix(doc, "Asimismo, aunque algunos estudios han abordado"),
    "El vacío no consiste en demostrar una optimización general de inventarios, rentabilidad o compras por volumen, pues esas decisiones requerirían recetas, existencias, costos de faltante y tiempos de entrega. El vacío se delimita a evaluar si los modelos de aprendizaje automático aportan una mejora verificable de precisión para anticipar el presupuesto de abastecimiento frente a referencias reproducibles.",
)
replace_text(
    find_prefix(doc, "En este sentido, se identifica un vacío de investigación"),
    "En este sentido, la presente investigación contribuye mediante un estudio de caso aplicado en Cup&Cake, microempresa de repostería creativa en Tizayuca, Hidalgo. Integra registros históricos de ventas y compras, usa validación Rolling-Origin y compara una línea base empírica, modelos estadísticos y modelos de aprendizaje automático para predecir el importe diario de compras. Su contribución se limita de forma explícita al apoyo del presupuesto de abastecimiento, sin inferir automáticamente una reducción de merma, inventario o rentabilidad.",
)

# Referencias APA 7 verificadas. Se corrige la ficha de Spiliotis ya existente y se
# elimina la duplicidad de M5 que tenía años y DOI inconsistentes.
replace_text(
    find_prefix(doc, "Spiliotis, E., Makridakis, S., Semenoglou, A. A., & Assimakopoulos, V."),
    "Spiliotis, E., Makridakis, S., Semenoglou, A. A., & Assimakopoulos, V. (2022). Comparison of statistical and machine learning methods for daily SKU demand forecasting. Operational Research, 22(3), 3037–3061. https://doi.org/10.1007/s12351-020-00605-2",
)
for duplicate in [
    "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2020). The M5 accuracy competition: Results, findings and conclusions. International Journal of Forecasting, 38(4), 1346–1364.",
    "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2021). The M5 accuracy competition: Results, findings and conclusions. International Journal of Forecasting, 38(4), 1346–1364. https://doi.org/10.1016/j.ijforecast.2021.11.013",
]:
    delete_paragraph(find_exact(doc, duplicate))

# Añadir las cuatro referencias nuevas, conservando el orden alfabético práctico.
anchor = find_prefix(doc, "García-Sánchez, A.")
insert_reference_after(doc, anchor, [
    ("Giannopoulos, P. G., Dasaklis, T. K., Tsantilis, I., & Patsakis, C. (2025). Machine learning algorithms in intermittent demand forecasting: A review. ", False),
    ("International Journal of Production Research", True),
    (", 1–43. https://doi.org/10.1080/00207543.2025.2578701", False),
])
anchor = find_prefix(doc, "Hastie, T.")
insert_reference_after(doc, anchor, [
    ("Hewamalage, H., Bergmeir, C., & Bandara, K. (2021). Recurrent neural networks for time series forecasting: Current status and future directions. ", False),
    ("International Journal of Forecasting, 37", True),
    ("(1), 388–427. https://doi.org/10.1016/j.ijforecast.2020.06.008", False),
])
anchor = find_prefix(doc, "Petropoulos, F.")
insert_reference_after(doc, anchor, [
    ("Poveda-Valverde, F., & Fierro Barragán, S. E. (2026). AI applications that can support sustainable practices in small and medium-sized enterprises in Latin America: A systematic review. ", False),
    ("Sustainability, 18", True),
    ("(7), Article 3603. https://doi.org/10.3390/su18073603", False),
])

doc.core_properties.title = "Predicción del presupuesto de abastecimiento mediante aprendizaje automático"
doc.save(OUTPUT)
print(OUTPUT)
