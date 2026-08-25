from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev36_(ZUJ)_22ago2026_EstadoDelArteActualizado.docx")
OUTPUT = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev37_(ZUJ)_22ago2026_Tabla1Reconstruida.docx")

ROWS = [
    ("Hübner et al. (2024)", "Panaderías; ventas, producción y desperdicio", "Caso cuantitativo con ML", "El pronóstico de demanda puede apoyar la reducción de desperdicio.", "Caso con infraestructura superior a la microempresa.", "Aporta el antecedente sectorial; justifica evaluar una solución acotada a registros internos."),
    ("Huber y Stuckenschmidt (2020)", "Retail diario; ventas y calendario", "Comparación de ML con variables calendáricas", "Las variables de calendario mejoran pronósticos en eventos.", "Contexto minorista de mayor escala.", "Sustenta incluir calendario y eventos disponibles en el dataset."),
    ("Abrar et al. (2024)", "Cadena de suministro; demanda multifuente", "Modelo híbrido CNN-LSTM con XAI", "Los modelos híbridos pueden integrar señales complejas.", "Alta complejidad y requisitos de datos.", "Sirve como contraste: la tesis evalúa alternativas parsimoniosas en Small Data."),
    ("Spiliotis et al. (2022)", "SKU diarios; series irregulares e intermitentes", "Comparación de métodos estadísticos y ML", "La superioridad de ML depende del objetivo y debe verificarse empíricamente.", "Datos minoristas con múltiples SKU.", "Fundamenta comparar Random Forest, ARIMA y línea base empírica."),
    ("Makridakis et al. (2018)", "Competencia M4; múltiples series y horizontes", "Benchmark de métodos estadísticos, híbridos y ML", "Los modelos complejos no dominan de forma universal.", "No estudia microempresas ni compras.", "Sustenta una hipótesis condicional y la evaluación comparativa."),
    ("Hewamalage et al. (2021)", "Series temporales con RNN", "Estudio empírico frente a ETS y ARIMA", "Las RNN son competitivas, pero no son soluciones universales.", "No se centra en presupuesto de compras de microempresas.", "Justifica reportar RNN/LSTM sin asumir superioridad frente a modelos simples."),
    ("Giannopoulos et al. (2025)", "Demanda intermitente en distintas industrias", "Revisión de ML y protocolos de evaluación", "La intermitencia exige atención a partición temporal, ajuste y métricas.", "Predomina evidencia de SKU y refacciones.", "Sustenta Rolling-Origin y el uso cauteloso de MAPE."),
    ("Borsato (2023)", "Repostería creativa y valor estético", "Estudio sectorial cualitativo", "La personalización distingue a la repostería creativa.", "No evalúa pronósticos ni IA.", "Delimita el contexto de negocio del caso Cup&Cake."),
    ("Ab Karim et al. (2020)", "Pastelería; creatividad y experiencia", "Estudio de retrato/teoría integrada", "La creatividad y personalización influyen en la oferta del sector.", "No modela demanda.", "Apoya la explicación de la variabilidad comercial del caso."),
    ("Ben Abdallah et al. (2022)", "Empresas pasteleras; innovación de producto", "Estudio exploratorio", "La innovación favorece desempeño y posicionamiento competitivo.", "No analiza pronóstico ni ML.", "Complementa la caracterización del sector de repostería."),
    ("Mikalef et al. (2019)", "Organizaciones orientadas a datos", "Modelo de capacidades analíticas", "Las capacidades analíticas se relacionan con innovación.", "No se limita a microempresas ni pronóstico.", "Sustenta la adopción de datos como capacidad organizacional."),
    ("Dubey et al. (2021)", "Manufactura; analítica, IA y desempeño", "Modelo empírico de desempeño operativo", "El valor de IA depende de capacidades y contexto.", "Contexto industrial.", "Refuerza que el DSS es apoyo y no garantía automática de desempeño."),
    ("Aprodu et al. (2025)", "Sector panadero e Industria 4.0", "Revisión sistemática", "La digitalización integra datos y automatización en panadería.", "Evidencia sectorial heterogénea.", "Apoya la pertinencia de digitalizar ventas y compras del caso."),
    ("Alekseeva et al. (2021)", "PYMES europeas; adopción de IA", "Estudio sobre barreras de adopción", "Recursos, infraestructura y talento limitan la adopción.", "No mide pronóstico de compras.", "Justifica una solución de bajo requerimiento basada en datos existentes."),
    ("Maldonado-Guzmán et al. (2018)", "PYMES mexicanas; innovación y desempeño", "Estudio empírico cuantitativo", "La adopción tecnológica se asocia con competitividad.", "No evalúa IA ni pronósticos.", "Aporta contexto nacional de transformación tecnológica."),
    ("Cuevas-Vargas et al. (2016)", "PYMES mexicanas; TIC y desempeño", "Estudio cuantitativo", "Las TIC favorecen la innovación y el desempeño organizacional.", "No incorpora ML.", "Sustenta el antecedente mexicano de digitalización."),
    ("Maldonado-Guzmán y Garza-Reyes (2020)", "PYMES mexicanas; transformación digital", "Estudio de procesos de digitalización", "Las tecnologías basadas en datos fortalecen decisiones.", "No compara modelos predictivos.", "Contextualiza la brecha entre digitalización y analítica predictiva."),
    ("Saavedra-García y Tapia-Sánchez (2020)", "PYMES mexicanas; adopción de TIC", "Análisis de adopción tecnológica", "Las restricciones financieras y de talento limitan la adopción.", "Enfoque descriptivo.", "Respalda la pertinencia de una solución reproducible y accesible."),
    ("Romero-Hidalgo et al. (2021)", "Pequeñas empresas; conocimiento e innovación", "Modelo de gestión del conocimiento", "El uso de datos favorece la innovación organizacional.", "No evalúa precisión predictiva.", "Complementa el valor de formalizar registros dispersos."),
    ("Poveda-Valverde y Fierro Barragán (2026)", "PYMES latinoamericanas; adopción de IA", "Revisión sistemática", "Persisten barreras de infraestructura, presupuesto y talento; logística concentra aplicaciones.", "No se centra en microempresas de repostería.", "Refuerza el vacío regional y la necesidad de evidencia aplicada en microempresas."),
]


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for line_index, line in enumerate(text.split("\n")):
        if line_index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(7.5 if not bold else 7.7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def set_width(cell, inches):
    width = int(inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(width))
    tc_width.set(qn("w:type"), "dxa")


doc = Document(SOURCE)
old_table = next(table for table in doc.tables if table.cell(0, 0).text.strip() == "Num")

headers = [
    "Autor(es) y año",
    "Contexto y datos",
    "Metodología",
    "Hallazgo relevante",
    "Limitación o brecha",
    "Relación con la investigación",
]
widths = [0.88, 1.0, 0.92, 1.18, 1.02, 1.2]
new_table = doc.add_table(rows=1, cols=len(headers))
new_table.style = "Table Grid"
new_table.autofit = False
new_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

for index, header in enumerate(headers):
    set_cell_text(new_table.rows[0].cells[index], header, bold=True, fill="D9EAF7")
    set_width(new_table.rows[0].cells[index], widths[index])

for entry in ROWS:
    cells = new_table.add_row().cells
    for index, value in enumerate(entry):
        set_cell_text(cells[index], value)
        set_width(cells[index], widths[index])

# Repite la fila de encabezados al continuar en páginas posteriores.
tr_pr = new_table.rows[0]._tr.get_or_add_trPr()
tbl_header = OxmlElement("w:tblHeader")
tbl_header.set(qn("w:val"), "true")
tr_pr.append(tbl_header)

# Sustituye la tabla antigua sin alterar la leyenda ni el párrafo de fuente.
old_table._element.addprevious(new_table._tbl)
old_table._element.getparent().remove(old_table._element)

caption = next(p for p in doc.paragraphs if p.text.startswith("Tabla 1 - Investigaciones relevantes"))
caption_text = "Tabla 1 - Estudios citados en el estado del arte y relación con la presente investigación"
caption.clear()
caption.add_run(caption_text)

source = next(p for p in doc.paragraphs if p.text.startswith("Fuente: Elaboración propia con base"))
source.clear()
source.add_run("Fuente: Elaboración propia con base en las referencias citadas en el Estado del Arte.")

doc.core_properties.title = "Predicción del presupuesto de abastecimiento mediante aprendizaje automático"
doc.save(OUTPUT)
print(OUTPUT)
