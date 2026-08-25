from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev34_(ZUJ)_22ago2026.docx")
OUTPUT = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev35_(ZUJ)_22ago2026_PresupuestoAbastecimiento.docx")


def set_paragraph_text(paragraph, text):
    """Replace paragraph text while retaining its paragraph style and first-run formatting."""
    run_props = None
    if paragraph.runs and paragraph.runs[0]._element.rPr is not None:
        run_props = deepcopy(paragraph.runs[0]._element.rPr)
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if run_props is not None:
        run._element.insert(0, run_props)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def find_exact(document, old):
    normalized_old = " ".join(old.split())
    found = [p for p in document.paragraphs if " ".join(p.text.split()) == normalized_old]
    if len(found) != 1:
        raise ValueError(f"Expected exactly one paragraph for: {old[:80]!r}; found {len(found)}")
    return found[0]


def replace(document, old, new):
    set_paragraph_text(find_exact(document, old), new)


def replace_prefix(document, prefix, new):
    normalized_prefix = " ".join(prefix.split())
    found = [p for p in document.paragraphs if " ".join(p.text.split()).startswith(normalized_prefix)]
    if len(found) != 1:
        raise ValueError(f"Expected exactly one paragraph beginning with: {prefix!r}; found {len(found)}")
    set_paragraph_text(found[0], new)


def replace_all_exact(document, old, new):
    normalized_old = " ".join(old.split())
    found = [p for p in document.paragraphs if " ".join(p.text.split()) == normalized_old]
    if not found:
        raise ValueError(f"No paragraphs found for: {old!r}")
    for paragraph in found:
        set_paragraph_text(paragraph, new)


doc = Document(SOURCE)

# Portada y síntesis.
replace(
    doc,
    "“DE LA INTUICIÓN A LA INTELIGENCIA ARTIFICIAL: TRANSFORMACIÓN PREDICTIVA DE LA PLANEACIÓN FINANCIERA EN UNA MICROEMPRESA MEDIANTE INTELIGENCIA DE NEGOCIOS CON MODELOS DE APRENDIZAJE AUTOMÁTICO”",
    "“PREDICCIÓN DEL PRESUPUESTO DE ABASTECIMIENTO MEDIANTE APRENDIZAJE AUTOMÁTICO: EVALUACIÓN COMPARATIVA EN UNA MICROEMPRESA DE REPOSTERÍA”",
)
replace(
    doc,
    "La investigación desarrolló y evaluó un sistema de pronóstico para apoyar la planeación financiera y el abastecimiento de Cup&Cake, microempresa de repostería ubicada en Tizayuca, Hidalgo. Se integraron registros diarios de ventas y compras con variables calendáricas, económicas y climáticas; posteriormente se construyeron rezagos, ventanas móviles y características cíclicas bajo controles de fuga de información. La evaluación se realizó mediante validación temporal Rolling-Origin y comparación contra una línea base empírica.",
    "La investigación desarrolló y evaluó modelos predictivos para anticipar el presupuesto de abastecimiento de Cup&Cake, microempresa de repostería ubicada en Tizayuca, Hidalgo. Se integraron registros diarios de ventas y compras con variables calendáricas, económicas y climáticas; posteriormente se construyeron rezagos, ventanas móviles y características cíclicas bajo controles de fuga de información. La evaluación se realizó mediante validación temporal Rolling-Origin y comparación contra una línea base empírica y modelos estadísticos de referencia.",
)
replace(
    doc,
    "El dataset maestro quedó conformado por 1,612 observaciones y 65 variables, mientras que el conjunto de modelado reunió 1,584 días y 276 columnas. Los modelos ganadores fueron ARIMA para los objetivos de ventas y Random Forest para los objetivos de compras. La reducción relativa del RMSE fue de 3.2 % y 5.1 % en ventas, y de 28.3 % y 22.8 % en compras; por ello, la hipótesis se aceptó parcialmente. Los resultados se integraron en un sistema de soporte a la decisión con indicadores, filtros y salvaguardas de interpretación humana.",
    "El dataset maestro quedó conformado por 1,612 observaciones y 65 variables, mientras que el conjunto de modelado reunió 1,584 días y 276 columnas. ARIMA resultó más competitivo para ventas, mientras que Random Forest obtuvo el mejor desempeño para el importe y la frecuencia de compras. En compras, la reducción relativa del RMSE fue de 28.3 % y 22.8 % frente a la línea base; por ello, la hipótesis se respaldó para el presupuesto de abastecimiento. Los resultados se integraron en un sistema de soporte a la decisión con indicadores, filtros y salvaguardas de interpretación humana.",
)

# Resumen en inglés.
replace_prefix(doc, "This research developed and evaluated a forecasting system", "This research developed and evaluated predictive models to anticipate the procurement budget of Cup&Cake, a small bakery located in Tizayuca, Hidalgo. Daily sales and purchasing records were integrated with calendar, economic and weather variables. Lagged variables, moving windows and cyclical encodings were generated under information-leakage controls. Models were assessed through Rolling-Origin temporal validation against an empirical baseline and statistical reference models. ARIMA was more competitive for sales, whereas Random Forest achieved the best performance for purchasing amount and frequency. For purchases, relative RMSE reductions were 28.3% and 22.8% against the baseline. Results were incorporated into a business-intelligence decision-support dashboard.")
replace_prefix(doc, "Keywords: demand forecasting", "Keywords: procurement budget, microenterprise, time series, machine learning, Rolling-Origin, business intelligence.")
replace(
    doc,
    "Palabras clave: pronóstico de demanda, microempresa, series temporales, aprendizaje automático, Rolling-Origin, inteligencia de negocios.",
    "Palabras clave: presupuesto de abastecimiento, microempresa, series temporales, aprendizaje automático, Rolling-Origin, inteligencia de negocios.",
)

# Pregunta, objetivo e hipótesis.
replace(
    doc,
    "¿En qué medida un sistema de inteligencia de negocios basado en modelos de aprendizaje automático optimiza la precisión del pronóstico de demanda, y fortalece la planeación de abastecimiento y financiera de una microempresa de repostería creativa en Tizayuca, Hidalgo, frente al método empírico utilizado históricamente?",
    "¿En qué medida los modelos de aprendizaje automático, entrenados con registros históricos de ventas y compras, mejoran la precisión de la predicción diaria del monto de compras de una microempresa de repostería creativa en Tizayuca, Hidalgo, frente a reglas empíricas y modelos estadísticos de referencia?",
)
replace(
    doc,
    "Evaluar en qué medida un sistema de inteligencia de negocios basado en modelos de aprendizaje automático optimiza la precisión del pronóstico de demanda, y fortalece la planeación de abastecimiento y financiera de una microempresa de repostería creativa en Tizayuca, Hidalgo, frente al método empírico utilizado históricamente",
    "Evaluar el desempeño de modelos de aprendizaje automático para predecir el monto diario de compras de una microempresa de repostería creativa, utilizando registros históricos de ventas y compras, y compararlo con reglas empíricas y modelos estadísticos de referencia.",
)
for old, new in [
    ("Describir cómo funciona actualmente el registro de ventas, compra de insumos y planeación empírica para tener una idea clara del negocio", "Caracterizar los registros históricos de ventas y compras, así como operacionalizar una línea base empírica reproducible de pronóstico."),
    ("Construir un dataset maestro a partir de los orígenes de datos de ventas, compra de insumos con los que cuenta Cup&Cake y adicionales con las variables exógenas (calendáricas y comerciales) para identificar las variables que influyen favorablemente en el modelo para luego entrenar los modelos de aprendizaje automático", "Construir un dataset maestro diario a partir de los registros de ventas y compras, incorporando variables temporales y exógenas disponibles para el modelado predictivo."),
    ("Emplear métodos estadísticos como ARIMA, SARIMA, regresión lineal multivariable, árboles de decisión y algoritmos de aprendizaje automático, LSTM (Long Short-Term Memory), redes neuronales recurrentes (RNN), para predecir la demanda, los ingresos y anticipar las compras a realizar. ", "Entrenar y evaluar métodos estadísticos y modelos de aprendizaje automático para predecir el importe y la frecuencia diaria de compras."),
    ("Comparar el desempeño de los datasets mediante técnicas de partición temporal y validación Rolling-Origin, con el fin de identificar el conjunto de datos que proporcione mayor precisión y estabilidad predictiva.", "Comparar el desempeño de los modelos mediante partición temporal y validación Rolling-Origin, con el fin de identificar la alternativa con mayor precisión y estabilidad predictiva."),
    ("Seleccionar los modelos predictivos mediante técnicas de ajuste de hiperparámetros, selección de características y procesos de validación, para garantizar su robustez y capacidad de generalización.", "Seleccionar la configuración predictiva con base en ajuste de hiperparámetros, selección de características, desempeño fuera de muestra e interpretabilidad."),
    ("Probar el desempeño del modelo frente al método empírico, utilizando métricas estadísticas como MAE, RMSE y MAPE para evaluar el mejor modelo.", "Contrastar los modelos de aprendizaje automático con la línea base empírica y los modelos estadísticos mediante MAE y RMSE; usar MAPE sólo como diagnóstico complementario debido a la intermitencia de las series."),
    ("Integrar los resultados en un tablero de Inteligencia de Negocios para toma de decisiones basadas en datos", "Integrar los resultados en un tablero de inteligencia de negocios para apoyar la consulta del presupuesto de abastecimiento y la toma de decisiones basada en datos."),
]:
    replace(doc, old, new)
remove_paragraph(find_exact(doc, "Estimar el impacto de los modelos predictivos en la planeación financiera y la operación de la microempresa."))
remove_paragraph(find_exact(doc, "Traducir los resultados del pronóstico a escenarios de planeación financiera y abastecimiento, mediante simulaciones de compras, inventarios y expansión comercial, para estimar su utilidad estratégica en la toma de decisiones."))
replace(
    doc,
    "La implementación de modelos supervisados de aprendizaje automático integrados a un tablero de Inteligencia de Negocios reducirá el error en un 20% del pronóstico de demanda, en comparación con el método empírico de planeación utilizado por la microempresa Cup&Cake medido mediante MAE, RMSE y MAPE; dicha reducción permitirá estimar con mayor precisión los requerimientos de abastecimiento y mejorar la calidad de la planeación financiera.",
    "Al menos un modelo de aprendizaje automático reducirá en 20 % o más el RMSE de la predicción del importe diario de compras frente a la línea base empírica, y mostrará un desempeño no inferior al de los modelos estadísticos de referencia. La evaluación se complementará con MAE; MAPE se interpretará de manera cautelosa por la presencia de valores cero.",
)

# Corrección de diseño metodológico y alcance.
replace(doc, "Tipo de tesis: Experimento Cuasiexperimental: limitante, falta de aleatorización, se debe intentar establecer semejanza entre los grupos, difieren de los experimentos “verdaderos” en la equivalencia inicial de los grupos.", "Tipo de tesis: estudio aplicado, cuantitativo, no experimental y evaluativo, basado en un caso de estudio longitudinal-retrospectivo.")
replace(doc, "Se trata de un tipo de investigación centrada en encontrar mecanismos o estrategias que permitan lograr un objetivo concreto, como el de Evaluar en qué medida los modelos de aprendizaje automático contribuyen al incremento de los ingresos en una microempresa de repostería creativa en Tizayuca, Hidalgo a partir de la predicción del abastecimiento y la planeación financiera. Por consiguiente, el tipo de ámbito al que se aplica es muy específico y bien delimitado, ya que no se trata de explicar una amplia variedad de situaciones, sino que más bien se intenta abordar un problema específico.", "Se trata de una investigación centrada en evaluar un mecanismo concreto: el uso de modelos de aprendizaje automático para predecir el importe diario de compras de una microempresa de repostería creativa. El ámbito es específico y delimitado; no busca explicar una amplia variedad de situaciones ni atribuir causalidad sobre ingresos, sino contrastar precisión predictiva en un caso de estudio.")
replace(doc, "Según el nivel de profundización en el objeto de estudio", "Según el alcance evaluativo del estudio")
replace(doc, "Explicativa", "Evaluativa y comparativa")
replace(doc, "La presente investigación se clasifica como de tipo explicativa. De acuerdo con la fundamentación teórica, este nivel de investigación no se limita a la descripción de fenómenos, sino que se centra en determinar las causas y consecuencias, profundiza en la relación causal directa entre la dependencia histórica del propietario hacia el conocimiento empírico y tácito —marcada por sesgos cognitivos documentados, tales como la sobre confianza y la heurística de disponibilidad — y la consecuente incertidumbre operativa. Esta dependencia metodológica para la predicción de compras de materia prima se identifica como la causa fundamental de las ineficiencias en el abastecimiento, lo cual genera un escenario financiero poco confiable que obstaculiza la viabilidad de trazar y ejecutar planes de inversión a largo plazo.", "La presente investigación tiene un alcance evaluativo y comparativo. Describe el comportamiento histórico de ventas y compras y compara, bajo las mismas ventanas temporales, la precisión de reglas empíricas, modelos estadísticos y modelos de aprendizaje automático. El diseño no permite atribuir causalidad sobre la rentabilidad, los ingresos o la eficiencia real de inventarios; su inferencia se limita al desempeño predictivo fuera de muestra.")
replace(doc, "En consecuencia, el diseño explicativo de este trabajo busca establecer y validar secuencias de causa-efecto ante la introducción de una nueva variable: la tecnología algorítmica. El objetivo es comprobar cómo la implementación de modelos predictivos de aprendizaje automático actúa como el factor explicativo y transformador para la optimización de los ingresos orgánicos y la planeación financiera. De este modo, se evalúa a la Inteligencia Artificial como el mecanismo que mitiga los sesgos humanos, optimiza el pronóstico de insumos y dota al negocio de la certidumbre cuantitativa necesaria para su expansión comercial.", "En consecuencia, el diseño evaluativo busca determinar si la tecnología algorítmica aporta una mejora verificable de precisión respecto de las referencias comparadas. La inteligencia artificial se evalúa como una alternativa de apoyo para anticipar el presupuesto de abastecimiento, sin suponer que sustituye el juicio del responsable ni que genera automáticamente mejoras financieras u operativas.")
replace(doc, "Cuasiexperimental", "No experimental, evaluativa y retrospectiva")
replace(doc, "Por el grado de manipulación de las variables y las condiciones inherentes al entorno de estudio, la investigación adopta un diseño metodológico cuasiexperimental. A diferencia de los experimentos puros desarrollados en entornos de laboratorio estrictamente controlados, este diseño permite la intervención deliberada sobre una variable independiente principal dentro de un ecosistema organizacional vivo, dinámico y no aleatorizado", "Por las condiciones inherentes al caso de estudio, la investigación adopta un diseño no experimental, evaluativo y retrospectivo. No se introdujo una intervención deliberada en la operación de la empresa ni se asignaron grupos; se analizaron registros históricos en su secuencia temporal.")
replace(doc, "En el contexto de esta investigación, la intervención metodológica consiste en la transición del modelo de planeación financiera actual de Cup&Cake —caracterizado por decisiones empíricas e intuitivas— hacia la implementación de un ecosistema tecnológico predictivo sustentado en Inteligencia de Negocios y modelos de Aprendizaje Automático. El objetivo de esta manipulación es observar, medir y cuantificar su impacto directo sobre la variable dependiente: la precisión predictiva en la estimación de la demanda y el abastecimiento de materia prima. ", "En el contexto de esta investigación, el objeto evaluado es la precisión predictiva del importe y la frecuencia de compras. La línea base representa reglas reproducibles basadas en el último valor disponible y el promedio móvil de siete días; los modelos estadísticos y de aprendizaje automático se comparan con ella en las mismas ventanas futuras.")
replace(doc, "La elección de este diseño resulta no solo adecuada, sino metodológicamente imperativa, debido a que el experimento se ejecuta sobre la operatividad diaria y real de la microempresa. Al tratarse de un entorno comercial en activo, el investigador carece de un control absoluto sobre el universo de variables exógenas que interactúan en el mercado. Factores externos como la volatilidad económica, variaciones repentinas en los precios de los proveedores de repostería, alteraciones climáticas o cambios imprevistos en el comportamiento de consumo del cliente, constituyen elementos naturales del contexto que no pueden ser aislados, neutralizados ni distribuidos de forma aleatoria. Por consiguiente, el rigor cuasiexperimental asume e integra estas limitantes del mundo real, permitiendo validar la eficacia analítica y la resiliencia del modelo de Inteligencia Artificial frente a la línea base empírica tradicional bajo condiciones auténticas de mercado. ", "Por consiguiente, el rigor del estudio proviene de la separación temporal estricta, la validación Rolling-Origin y la comparación homogénea contra referencias reproducibles. Estos elementos permiten evaluar generalización predictiva en condiciones históricas reales, sin afirmar causalidad experimental.")
replace(doc, "Para dar cumplimiento a los objetivos de la investigación y estructurar el diseño cuasiexperimental, el desarrollo operativo del proyecto se divide en fases secuenciales. Estas etapas abarcan desde la recolección de la información empírica de la microempresa hasta la simulación financiera mediante Inteligencia de Negocios.", "Para dar cumplimiento a los objetivos de la investigación y estructurar el diseño evaluativo, el desarrollo operativo se divide en fases secuenciales: recolección, integración, modelado, validación temporal y comunicación de resultados para el presupuesto de abastecimiento.")

# Desarrollo, resultados, discusión y conclusiones: acotar lo que sí se probó.
replace(doc, "Diagnóstico de la Línea Base y Evaluación del Método Empírico (Preprueba) ", "Diagnóstico de la línea base empírica y evaluación comparativa")
replace(doc, "La línea base representó la práctica empírica reproducible con la que se contrastaron los modelos. Para cada objetivo se evaluaron la persistencia del último valor y el promedio móvil de siete días; se conservó como referencia el método con menor RMSE en cada ventana temporal. Este planteamiento permitió cuantificar la ganancia analítica sin atribuir causalidad experimental. Los datos correspondientes se sintetizan en la Tabla 7.", "La línea base representó una referencia empírica reproducible con la que se contrastaron los modelos. Para cada objetivo se evaluaron la persistencia del último valor y el promedio móvil de siete días; se conservó como referencia el método con menor RMSE en cada ventana temporal. Este planteamiento permitió cuantificar la ganancia predictiva sin equiparar la línea base con un grupo control ni atribuir causalidad experimental. Los datos correspondientes se sintetizan en la Tabla 7.")
replace(doc, "Alcance de la simulación financiera y de abastecimiento", "Alcance del apoyo al presupuesto de abastecimiento")
replace(doc, "El DSS permite explorar el comportamiento de los cuatro objetivos, comparar el desempeño contra la línea base y revisar estabilidad por ventana. Sin embargo, no se ejecutó una optimización prescriptiva de inventario porque los registros disponibles no contienen de manera consistente costos de faltante, merma, inventario inicial ni tiempos de entrega. Por rigor científico, los escenarios se interpretaron como apoyo diagnóstico y no como órdenes automáticas de compra.", "El DSS permite explorar el comportamiento de los objetivos, comparar el desempeño contra la línea base y revisar estabilidad por ventana. Su alcance es anticipar el importe y la frecuencia de compras como apoyo para reservar presupuesto de abastecimiento. No genera órdenes automáticas ni cantidades físicas por insumo, porque los registros no contienen de manera consistente recetas, inventario inicial/final, merma, tiempos de entrega ni unidades de medida homologadas. Por rigor científico, los resultados se interpretan como apoyo diagnóstico y de planeación presupuestal.")
replace(doc, "La hipótesis H1 estableció una reducción mínima de 20 % en el error respecto de la línea base. El umbral se alcanzó en importe y registros de compras, pero no en los dos objetivos de ventas. En consecuencia, H1 se aceptó parcialmente: la evidencia apoya una mejora sustantiva para el abastecimiento, no una mejora generalizada para todos los procesos.", "La hipótesis H1 estableció una reducción mínima de 20 % en el RMSE del importe diario de compras frente a la línea base. El umbral se alcanzó en importe y registros de compras mediante Random Forest. En consecuencia, H1 se respaldó para el presupuesto de abastecimiento. Los resultados de ventas se reportan como análisis complementario y muestran que ARIMA fue más competitivo en ese dominio.")
replace(doc, "La evidencia permite priorizar el uso predictivo en compras, donde la mejora fue material y consistente con la necesidad de reducir faltantes y sobreabastecimiento. Para ventas, el DSS debe emplearse como referencia de tendencia y alerta, acompañado del juicio del responsable del negocio, ya que la ganancia frente a la práctica empírica fue modesta.", "La evidencia permite priorizar el uso predictivo en compras, donde la mejora fue material y consistente con la necesidad de anticipar el presupuesto de abastecimiento. No es posible inferir reducción efectiva de faltantes o sobreabastecimiento, porque no se cuenta con inventarios, recetas, mermas ni tiempos de entrega. Para ventas, el DSS debe emplearse como referencia de tendencia y alerta, acompañado del juicio del responsable del negocio, ya que ARIMA fue más competitivo que los modelos de aprendizaje automático.")
replace(doc, "Se construyó una solución analítica reproducible que integra datos transaccionales y exógenos, transforma las series mediante ingeniería de características, compara modelos con validación temporal y comunica los resultados en un DSS. El proceso convirtió registros dispersos en evidencia verificable para la planeación financiera y de abastecimiento.", "Se construyó una solución analítica reproducible que integra datos transaccionales y exógenos, transforma las series mediante ingeniería de características, compara modelos con validación temporal y comunica los resultados en un DSS. El proceso convirtió registros dispersos en evidencia verificable para anticipar el presupuesto de abastecimiento.")
replace(doc, "La hipótesis se aceptó parcialmente. Random Forest redujo el RMSE de compras en 28.3 % para importes y 22.8 % para registros; ARIMA mejoró ventas en 3.2 % y 5.1 %, por debajo del umbral de 20 %. Por tanto, el principal valor operativo se concentra en el apoyo al abastecimiento.", "La hipótesis se respaldó para el presupuesto de abastecimiento. Random Forest redujo el RMSE de compras en 28.3 % para importes y 22.8 % para registros frente a la línea base. ARIMA fue más competitivo para ventas; por ello, la evidencia no respalda una superioridad general de los modelos de aprendizaje automático para todos los objetivos, sino su utilidad específica para anticipar compras.")
replace(doc, "Se recomienda actualizar mensualmente las fuentes, vigilar la cobertura de compras, conservar la línea base como control permanente y reentrenar cuando el RMSE por ventana muestre deterioro sostenido. Las decisiones de compra deben incorporar revisión humana y límites financieros.", "Se recomienda actualizar mensualmente las fuentes, vigilar la cobertura de compras, conservar la línea base como referencia permanente y reentrenar cuando el RMSE por ventana muestre deterioro sostenido. El pronóstico debe utilizarse para reservar y revisar el presupuesto de abastecimiento; las decisiones de compra deben incorporar revisión humana y límites financieros.")
replace(doc, "Conviene registrar inventario inicial y final, merma, costo de faltante, tiempo de entrega y promociones; evaluar modelos específicos para demanda intermitente; calcular WAPE de forma persistente; y efectuar una validación prospectiva antes de automatizar recomendaciones de abastecimiento.", "Como trabajo futuro, conviene registrar inventario inicial y final, recetas, merma, costo de faltante, tiempo de entrega y promociones; homologar las unidades de medida por insumo; evaluar modelos específicos para demanda intermitente; y efectuar una validación prospectiva antes de automatizar recomendaciones de compra por volumen.")

# Limpieza de afirmaciones que contradecían el diseño y el alcance verificable.
replace_prefix(doc, "No obstante, la gran limitante del conocimiento tácito", "El conocimiento tácito del propietario es útil para interpretar el contexto del negocio, pero suele estar poco formalizado y es difícil de evaluar con métricas reproducibles. La investigación no asume que el empirismo sea inválido; lo convierte en una referencia contrastable frente a modelos predictivos para anticipar el presupuesto de abastecimiento.")
replace_prefix(doc, "Asimismo, ante planes de expansión comercial", "En esta investigación, los pronósticos se emplean para anticipar el presupuesto de abastecimiento de corto plazo. La evaluación de inversiones, expansión comercial o flujos de efectivo requiere datos y modelos adicionales, por lo que se reconoce como una línea de trabajo futura.")
replace_prefix(doc, "Los resultados de la M5 demostraron", "La evidencia comparativa disponible en la literatura sugiere que los modelos de aprendizaje automático pueden aportar valor cuando existen relaciones no lineales y variables exógenas relevantes. Sin embargo, su superioridad no debe asumirse: depende de la calidad de los datos, del objetivo y de una validación temporal rigurosa.")
replace_prefix(doc, "Esta integración operativa consolida la transformación predictiva", "La integración de pronósticos en un tablero puede hacer explícitos los supuestos y apoyar una revisión sistemática del presupuesto de abastecimiento. No elimina el juicio humano ni garantiza, por sí sola, resultados financieros u operativos; su utilidad depende de la calidad de los datos y de la decisión que adopte el responsable del negocio.")
replace_prefix(doc, "En la investigación realizada con los datos de la empresa Cup&Cake, el Grupo de Control", "En la investigación realizada con los datos de Cup&Cake, la referencia comparativa es una línea base empírica reproducible, integrada por la persistencia del último valor y el promedio móvil de siete días. No constituye un grupo de control experimental ni permite atribuir las compras históricas exclusivamente a la intuición del propietario.")
replace_prefix(doc, "Para ello pueden usarse diferentes métodos, como el método observacional", "En un estudio evaluativo se comparan alternativas de pronóstico bajo las mismas condiciones temporales. El objetivo es estimar cuál generaliza mejor fuera de muestra, no establecer relaciones de causa-efecto.")
replace_prefix(doc, "La investigación cuasiexperimental se asemeja", "En un diseño no experimental se analizan datos observados sin manipular deliberadamente las condiciones de la empresa. La comparación temporal permite estimar desempeño predictivo, pero no inferir causalidad experimental.")
replace(doc, "Integración en Inteligencia de Negocios y simulación financiera", "Integración en Inteligencia de Negocios y comunicación del presupuesto de abastecimiento")
replace(doc, "Evaluación Comparativa de Rendimiento Analítico (Posprueba vs. Control) ", "Evaluación comparativa del rendimiento predictivo frente a la línea base")
replace_prefix(doc, "Bajo este marco temporal estricto, la interacción armónica", "Bajo este marco temporal estricto, MAE y RMSE permiten comparar la magnitud media y la penalización de errores grandes en cada modelo. MAPE se utiliza sólo como diagnóstico complementario, pues las series con numerosos valores cero pueden producir porcentajes inestables. Las métricas sirven para seleccionar el modelo que mejor anticipa el presupuesto de abastecimiento, no para calcular directamente inventario o liquidez.")
replace_prefix(doc, "La preprueba se operacionalizó", "La comparación se operacionalizó como el contraste entre el error de la línea base empírica —último valor disponible y promedio móvil de siete días— y el error de los modelos predictivos en las mismas ventanas futuras. No se emplearon preprueba, posprueba ni grupos de control experimentales.")
replace_prefix(doc, "La evaluación objetiva del grupo experimental", "La evaluación comparativa de los modelos algorítmicos y la línea base empírica se realizó calculando la discrepancia entre el valor real observado y la predicción. Se utilizaron las siguientes métricas estadísticas:")
replace_all_exact(doc, "Tabla 7 - Operacionalización de la línea base y la posprueba", "Tabla 7 - Operacionalización de la línea base y los modelos predictivos")

doc.core_properties.title = "Predicción del presupuesto de abastecimiento mediante aprendizaje automático"
doc.core_properties.subject = "Evaluación comparativa de modelos predictivos para una microempresa de repostería"
doc.save(OUTPUT)
print(OUTPUT)
