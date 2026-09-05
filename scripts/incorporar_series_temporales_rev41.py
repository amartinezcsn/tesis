"""Genera la revisión 41 incorporando el análisis temporal aprobado."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


PROJECT = Path(r"C:\Python\tesis")
SOURCE = PROJECT / "documentacion" / "TESIS_AGO2026_Rev40_(ZUJ)_24ago2026_HorizonteSemanal.docx"
OUTPUT = PROJECT / "documentacion" / "TESIS_AGO2026_Rev41_(ZUJ)_29ago2026_AnalisisSeriesTemporales.docx"
FIGURES = PROJECT / "output" / "semanal" / "series_temporales"
EXPECTED_TITLE = "PRONÓSTICO SEMANAL DEL PRESUPUESTO DE ABASTECIMIENTO MEDIANTE MODELOS ESTADÍSTICOS Y APRENDIZAJE AUTOMÁTICO CON VALIDACIÓN DE VENTANA DESLIZANTE: ESTUDIO DE CASO EN UNA MICROEMPRESA DE REPOSTERÍA CREATIVA DE TIZAYUCA, HIDALGO"


def find_contains(document: Document, fragment: str):
    for paragraph in document.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    raise ValueError(f"No se encontró el texto de anclaje: {fragment}")


def find_exact(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"No se encontró el párrafo exacto: {text}")


def replace(paragraph, text: str) -> None:
    paragraph.text = text


def add_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = paragraph._parent.add_paragraph()
    created._p.getparent().remove(created._p)
    new_p.addnext(created._p)
    created._p.getparent().remove(created._p)
    new_p.getparent().replace(new_p, created._p)
    if style:
        created.style = style
    if text:
        created.add_run(text)
    return created


def insert_sequence_after(anchor, items):
    current = anchor
    for style, text in items:
        current = add_after(current, text, style)
    return current


def insert_before(anchor, text: str = "", style: str | None = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def insert_picture_before(anchor, image_path: Path, caption: str):
    picture_paragraph = insert_before(anchor, style="Normal")
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(6.25))
    caption_paragraph = insert_before(anchor, caption, style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def enable_field_updates(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    document = Document(SOURCE)
    original_title = document.paragraphs[3].text.strip()
    if original_title != EXPECTED_TITLE:
        raise ValueError("El título de la revisión 40 no coincide con el título aprobado para conservar.")

    # Resumen y abstract: análisis temporal explícito y resultados corregidos.
    replace(find_contains(document, "La investigación evalúa modelos predictivos para pronosticar semanalmente"),
        "La investigación evalúa modelos estadísticos y de aprendizaje automático para pronosticar semanalmente el presupuesto de abastecimiento de Cup&Cake, microempresa de repostería ubicada en Tizayuca, Hidalgo. A partir de registros diarios de ventas y compras se construyó una serie semanal integrada con variables calendáricas, económicas y climáticas disponibles antes del periodo pronosticado. Antes del modelado se realizó un análisis explícito de cobertura, tendencia, estacionariedad, autocorrelación, intermitencia y valores atípicos, utilizando únicamente el bloque temporal continuo y reservando las últimas 16 semanas para evaluación.")
    replace(find_contains(document, "El estudio comparará modelos estadísticos y de aprendizaje automático"),
        "El análisis identificó una serie errática, con ADI de 1.16, CV² de 1.03 y 13.5% de semanas con importe cero dentro del bloque de desarrollo. La evaluación mediante ventana deslizante mostró que, para h=1, el ingenuo estacional de 52 semanas obtuvo el menor RMSE descriptivo (695.19), seguido por Random Forest con transformación log1p (699.28); para h=4, la mejor configuración descriptiva fue HistGradientBoosting tipo hurdle con variables exógenas (RMSE 399.56). Ningún modelo superó significativamente la línea base primaria después del ajuste Holm. Los resultados se comunican mediante un sistema de soporte a la decisión que conserva la interpretación humana y evita presentar la inteligencia artificial como sustituto automático del criterio empresarial.")
    replace(find_contains(document, "This research evaluates predictive models for weekly procurement-budget forecasting"),
        "This research evaluates statistical and machine-learning models for weekly procurement-budget forecasting at Cup&Cake, a small bakery in Tizayuca, Hidalgo. Daily sales and purchasing records were aggregated into a weekly series and complemented with calendar, economic, and weather variables available before each forecast. Prior to modeling, the study explicitly assessed coverage, trend, stationarity, autocorrelation, intermittency, and robust outliers using only the main continuous block while reserving the final 16 weeks for evaluation. The series was classified as erratic. At h=1, the 52-week seasonal naive benchmark achieved the lowest descriptive RMSE (695.19), closely followed by log-transformed Random Forest (699.28); at h=4, the best descriptive result was obtained by an exogenous hurdle HistGradientBoosting configuration (RMSE 399.56). No model significantly outperformed the primary baseline after Holm adjustment. Results support human review through a business-intelligence decision-support dashboard.")

    # Estado del arte.
    aggregation_anchor = find_contains(document, "El consolidado mensual se obtendrá a partir de los pronósticos semanales")
    insert_sequence_after(aggregation_anchor, [
        ("Heading 3", "Análisis de series temporales para demanda errática y Small Data"),
        ("Normal", "El pronóstico constituye una fase posterior a la caracterización de la serie. Antes de comparar algoritmos es necesario examinar la continuidad de la cobertura, la frecuencia de observación, la proporción de ceros, la tendencia, la autocorrelación y la estabilidad de los patrones estacionales. En Small Data, estos diagnósticos permiten limitar la complejidad del modelo y evitar que una arquitectura de inteligencia artificial aprenda discontinuidades producidas por ausencia de registro como si fueran comportamiento económico real (Hyndman & Athanasopoulos, 2021)."),
        ("Normal", "En series con semanas sin movimiento, el intervalo promedio entre ocurrencias (ADI) y el cuadrado del coeficiente de variación de los importes positivos (CV²) permiten distinguir comportamientos suaves, intermitentes, erráticos e irregulares. Esta clasificación no determina por sí sola el modelo ganador, pero fundamenta la inclusión de métodos especializados como Croston-SBA y TSB, así como transformaciones logarítmicas y modelos de dos etapas (Syntetos & Boylan, 2005)."),
        ("Normal", "El análisis temporal también debe separar la descripción retrospectiva de la selección predictiva. Las gráficas pueden mostrar toda la cobertura para fines de auditoría; sin embargo, las pruebas de estacionariedad, autocorrelación y selección de características que influyen en la configuración del modelo deben calcularse sin observar el bloque final de evaluación. Esta separación preserva el carácter fuera de muestra de la comparación entre métodos estadísticos y aprendizaje automático."),
    ])

    # Objetivo específico adicional.
    objective_anchor = find_contains(document, "Construir un dataset semanal a partir de los registros diarios")
    add_after(objective_anchor,
        "Caracterizar la estructura temporal de la serie semanal del importe de compras mediante el análisis de cobertura, tendencia, estacionariedad, autocorrelación, intermitencia, dispersión y posibles valores atípicos, con el propósito de fundamentar la selección y configuración de modelos estadísticos y de aprendizaje automático.",
        "Normal")

    # Marco teórico: sección nueva inmediatamente antes de modelos de pronóstico.
    model_heading = find_exact(document, "Modelos de pronóstico: Enfoques estadísticos y Aprendizaje Automático")
    theoretical_items = [
        ("Heading 2", "Análisis de series temporales como fundamento del modelado predictivo"),
        ("Normal", "Una serie temporal es una secuencia de observaciones ordenadas cuya posición cronológica contiene información. En este tipo de datos, el orden no puede ignorarse: el nivel, la tendencia, la estacionalidad, la dependencia serial y las perturbaciones excepcionales condicionan tanto la ingeniería de características como la capacidad de generalización de los modelos de inteligencia artificial (Box et al., 2015; Hyndman & Athanasopoulos, 2021)."),
        ("Heading 3", "Frecuencia, continuidad y cobertura"),
        ("Normal", "La frecuencia de análisis debe corresponder con la decisión empresarial. En esta investigación, la semana completa de lunes a domingo es la unidad del presupuesto de abastecimiento. La agregación reduce parte del ruido diario, pero disminuye el tamaño de muestra; por ello, antes de modelar se debe verificar la continuidad del calendario y distinguir un cero observado de una ausencia de captura. Una racha prolongada sin compras mientras continúan las ventas constituye evidencia de cobertura indeterminada y no debe tratarse automáticamente como demanda igual a cero."),
        ("Heading 3", "Tendencia, estacionalidad y estacionariedad"),
        ("Normal", "La tendencia representa cambios persistentes en el nivel de la serie, mientras que la estacionalidad describe patrones que se repiten con periodicidad relativamente estable. La descomposición STL permite visualizar de manera robusta los componentes de tendencia, estacionalidad y residuo; no obstante, una estacionalidad anual semanal requiere varios ciclos completos para ser interpretada con confianza. Cuando el bloque continuo contiene menos de tres ciclos de 52 semanas, la descomposición debe considerarse exploratoria y no una prueba concluyente."),
        ("Normal", "La estacionariedad supone estabilidad aproximada de la media, la varianza y la dependencia temporal. La prueba Dickey-Fuller aumentada (ADF) contrasta una raíz unitaria, mientras que KPSS parte de la hipótesis complementaria de estacionariedad. Su interpretación conjunta reduce el riesgo de decidir a partir de una sola prueba, aunque en muestras pequeñas ambas tienen potencia limitada. Las transformaciones log1p y la diferenciación pueden evaluarse para estabilizar la escala, siempre dentro de cada ventana de entrenamiento."),
        ("Heading 3", "Autocorrelación y diagnóstico de rezagos"),
        ("Normal", "La función de autocorrelación (ACF) mide la relación lineal entre la serie y sus valores rezagados; la función de autocorrelación parcial (PACF) aísla la asociación de cada rezago después de controlar los intermedios. Estos diagnósticos orientan la selección parsimoniosa de rezagos y órdenes autorregresivos, pero no autorizan ajustar la configuración con las semanas reservadas para evaluación. La prueba Ljung-Box complementa el análisis al examinar conjuntamente si persiste dependencia serial en varios rezagos."),
        ("Heading 3", "Intermitencia, variabilidad y valores atípicos"),
        ("Normal", "El ADI expresa el número promedio de periodos por ocurrencia positiva y CV² mide la variabilidad relativa de los importes distintos de cero. Con los umbrales de referencia ADI=1.32 y CV²=0.49, la combinación de baja intermitencia y alta variabilidad se clasifica como errática; valores altos en ambos indicadores corresponden a una serie irregular o lumpy (Syntetos & Boylan, 2005). Esta taxonomía justifica comparar referencias simples, modelos para demanda intermitente y algoritmos de aprendizaje automático robustos a ceros y picos."),
        ("Normal", "Los valores atípicos pueden representar errores de captura o eventos comerciales legítimos. Por ello, se identifican mediante mediana y desviación absoluta mediana (MAD), pero no se eliminan automáticamente. Su tratamiento requiere trazabilidad y conocimiento del negocio, pues los picos asociados con temporadas o celebraciones contienen información relevante para el aprendizaje supervisado."),
        ("Heading 3", "Relación con la Inteligencia Artificial"),
        ("Normal", "El análisis de series temporales no desplaza el enfoque de Inteligencia Artificial; establece las condiciones para emplearlo con rigor. Sus hallazgos definen rezagos, transformaciones, variables cíclicas, estrategias de regularización y modelos candidatos. De este modo, la IA se evalúa sobre representaciones coherentes con el tiempo y se compara con referencias reproducibles, evitando atribuir a la complejidad algorítmica una superioridad que sólo puede demostrarse fuera de muestra."),
    ]
    for style, text in theoretical_items:
        insert_before(model_heading, text, style)

    # Sustituye afirmaciones absolutas incompatibles con el diseño comparativo.
    replace(find_contains(document, "Intentar proyectar esta intrincada red de estacionalidades múltiples"),
        "La coexistencia de eventos calendáricos, variabilidad de montos y posibles relaciones no lineales justifica comparar métodos univariados con algoritmos de aprendizaje automático capaces de incorporar variables exógenas. La elección no puede resolverse de forma teórica: depende de la cobertura, el horizonte, la estabilidad temporal y el desempeño fuera de muestra bajo el mismo protocolo de evaluación.")
    replace(find_contains(document, "El modelo empírico de gestión, condicionado por los sesgos cognitivos humanos"),
        "El modelo empírico de gestión puede perder precisión ante variaciones rápidas, picos y dependencias temporales múltiples. Los modelos estadísticos y de aprendizaje automático ofrecen alternativas sistemáticas para representar esas señales; sin embargo, deben considerarse instrumentos de apoyo y sólo pueden declararse superiores cuando la comparación fuera de muestra demuestra una mejora estable y estadísticamente defendible.")

    # Metodología: fase diagnóstica antes de la ingeniería de características.
    methodology_anchor = find_contains(document, "Después de depurar y homologar los registros, las observaciones se agregarán por semana calendario")
    insert_sequence_after(methodology_anchor, [
        ("Heading 3", "Análisis exploratorio de la serie temporal semanal"),
        ("Normal", "Antes de generar características y comparar algoritmos se realizará un diagnóstico temporal reproducible de la serie objetivo. La cobertura se auditará con el calendario semanal y con la actividad de ventas. Las rachas prolongadas de compras iguales a cero se considerarán brechas de cobertura cuando coincidan con ventas positivas o formen una cola posterior al último registro; los ceros aislados dentro de un bloque continuo se conservarán como observaciones operativas."),
        ("Normal", "El bloque continuo principal se separará cronológicamente en desarrollo y evaluación. Las últimas 16 semanas quedarán reservadas. Las pruebas ADF y KPSS, las funciones ACF y PACF, Ljung-Box, ADI, CV², los estadísticos móviles y la detección robusta de atípicos se calcularán con el bloque de desarrollo. La descomposición STL con periodo 52 se mostrará sólo como recurso exploratorio cuando existan menos de tres ciclos anuales completos."),
        ("Normal", "Los diagnósticos no se utilizarán para aceptar H1 o H2. Su función será fundamentar la selección de rezagos, transformaciones y familias de modelos. Cualquier configuración derivada de estos resultados se ajustará exclusivamente con información anterior a cada origen de pronóstico, preservando la separación temporal y evitando fuga de información."),
        ("Normal", "La ejecución exportará tablas abiertas, figuras, un resumen JSON y el libro 01b_diagnostico_series_temporales.xlsx. El archivo documentará cobertura, estacionariedad, dependencia serial, intermitencia, atípicos y las limitaciones de interpretación, de modo que la caracterización pueda auditarse independientemente del tablero de inteligencia de negocios."),
    ])

    # Desarrollo: aclara que lo diario es una capa de integración, no el experimento principal.
    replace(find_contains(document, "El propósito de este trabajo fue construir una base de datos unificada, diaria y lista para modelado"),
        "El propósito de esta etapa fue construir una base diaria unificada y auditable como capa intermedia de integración. Esta base no constituye la unidad principal del experimento predictivo; su función es depurar las fuentes y permitir posteriormente la agregación semanal necesaria para el presupuesto de abastecimiento.")
    replace(find_contains(document, "Predecir la demanda o ventas diarias."), "Conservar el historial diario de ventas como fuente para construir predictores semanales.")
    replace(find_contains(document, "Utilizar esa predicción como base para estimar compras futuras."), "Agregar las compras observadas por semana calendario para formar la variable objetivo del presupuesto de abastecimiento.")
    replace(find_contains(document, "Integrar variables exógenas que ayuden al modelo a capturar estacionalidad"), "Integrar variables exógenas en su fecha de disponibilidad y trasladarlas posteriormente a la frecuencia semanal sin utilizar información futura.")
    replace(find_contains(document, "La salida final se preparó como un dataset maestro con una fila por día"),
        "La salida diaria intermedia contiene una fila por fecha entre 2022-01-01 y 2026-05-31. A partir de ella se construyó dataset_maestro_semanal.xlsx, con semanas completas de lunes a domingo. Sólo el bloque continuo con cobertura defendible de compras se utiliza en el análisis principal.")
    replace(find_contains(document, "Se generó un archivo final llamado dataset_maestro_diario.xlsx"),
        "Se generaron dataset_maestro_diario.xlsx como evidencia de integración y dataset_maestro_semanal.xlsx como base de la unidad analítica principal.")
    replace(find_contains(document, "Como resultado del proceso de depuración, homologación e integración, el dataset maestro final quedó conformado"),
        "Como resultado de la integración, el dataset maestro diario quedó conformado por 1,612 observaciones y 65 variables. Este conteo documenta la capa de origen; después de la agregación se obtuvieron 230 semanas calendario. La auditoría temporal identificó un bloque continuo principal de 127 semanas entre 2022-01-03 y 2024-06-03, una brecha interna de 44 semanas y una cola de 43 semanas sin cobertura confiable de compras.")
    replace(find_contains(document, "Las verificaciones de calidad confirmaron que no existen fechas omitidas"),
        "La continuidad del calendario diario no implica cobertura sustantiva de todas las fuentes. Aunque no existen fechas omitidas en el panel, una racha de compras igual a cero entre 2024-06-10 y 2025-04-07 coincide con ventas positivas en 77.3% de las semanas, por lo que se clasificó como brecha de registro y no como ausencia real de compras. Los ceros aislados dentro del bloque continuo se conservaron como actividad observada.")
    daily_feature_anchor = find_exact(document, "Ingeniería de Características (Feature Engineering)")
    add_after(daily_feature_anchor,
        "Las transformaciones diarias que se describen a continuación se conservan como antecedente de integración y trazabilidad. El experimento principal utiliza su versión semanal, con rezagos de 1, 2, 4, 8 y 52 semanas, ventanas móviles de 4, 8 y 12 semanas y variables exógenas controladas por disponibilidad.",
        "Normal")
    replace(find_contains(document, "Para el análisis se definieron como variables objetivo principales target_ventas_importe_real_2026_05"),
        "En la etapa diaria se conservaron variables monetarias de ventas y compras como insumos intermedios. Para la evaluación principal se definió una sola variable objetivo: target_compras_importe_semanal, expresada en pesos reales de mayo de 2026. Las ventas se utilizan únicamente como predictores históricos y como evidencia auxiliar para auditar la cobertura de compras.")
    replace(find_contains(document, "El procedimiento generó cuatro variables objetivo asociadas con el comportamiento diario"),
        "El pipeline heredado conserva cuatro objetivos diarios para trazabilidad; no forman parte del contraste principal. La investigación evalúa el importe semanal de compras, mientras que los registros de ventas y su importe se emplean como antecedentes históricos disponibles antes del pronóstico.")
    replace(find_contains(document, "Finalmente, el resultado se exportó al archivo dataset_modelado_diario.xlsx"),
        "El resultado diario se conservó en dataset_modelado_diario.xlsx como antecedente reproducible. El pipeline vigente genera dataset_modelado_semanal.xlsx, cuyo diccionario documenta la fuente, el rezago y la disponibilidad de cada predictor utilizado en los modelos estadísticos y de aprendizaje automático.")
    replace(find_contains(document, "El archivo dataset_modelado_diario.xlsx integra variables calendáricas y exógenas"),
        "El archivo dataset_modelado_diario.xlsx documenta la expansión dimensional heredada. Para responder la pregunta de investigación se utiliza dataset_modelado_semanal.xlsx, que integra variables calendáricas y exógenas, rezagos semanales, estadísticas móviles e indicadores de eventos. La selección se ejecuta dentro de cada ventana de entrenamiento y no sobre todo el historial.")

    # Nuevo bloque de desarrollo con evidencia generada.
    baseline_heading = find_exact(document, "Diagnóstico de la línea base empírica y evaluación comparativa")
    development_items = [
        ("Heading 2", "Análisis exploratorio de la serie temporal semanal"),
        ("Heading 3", "Auditoría de cobertura y selección del bloque continuo"),
        ("Normal", "El calendario semanal contenía 230 periodos. La auditoría identificó dos brechas prolongadas: 44 semanas entre 2024-06-10 y 2025-04-07, coincidentes con ventas positivas en 77.3% de las semanas, y una cola de 43 semanas a partir de 2025-08-04. Además, las 16 semanas posteriores a la brecha interna no aportaban las 52 semanas continuas requeridas para el entrenamiento. Por ello, el análisis principal seleccionó el bloque 2022-01-03 a 2024-06-03, integrado por 127 semanas; 111 se destinaron al diagnóstico y desarrollo y 16 se reservaron para evaluación."),
        ("Normal", "Esta decisión modifica la interpretación de los ceros. Los ceros aislados del bloque continuo representan semanas sin compras registradas y se conservan. En cambio, las rachas prolongadas con ventas activas se consideran ausencia de cobertura y se excluyen del análisis principal. Así se evita que los algoritmos de inteligencia artificial aprendan como patrón operativo una interrupción del proceso de captura."),
    ]
    for style, text in development_items:
        insert_before(baseline_heading, text, style)
    insert_picture_before(baseline_heading, FIGURES / "01_serie_semanal_compras.png", "Figura AT-1 - Serie semanal del importe de compras y bloque final reservado")
    for style, text in [
        ("Heading 3", "Estructura temporal, estacionariedad y dependencia serial"),
        ("Normal", "En el bloque de desarrollo, ADF no rechazó la presencia de raíz unitaria en nivel (p=0.200) ni después de la transformación log1p (p=0.078). KPSS rechazó la estacionariedad en ambos casos (p=0.010). La evidencia conjunta indica que la serie no puede tratarse como estacionaria en nivel y fundamenta comparar transformaciones, modelos con componentes dinámicos y algoritmos regularizados."),
        ("Normal", "La ACF presentó dependencia positiva en el primer rezago (0.217), mientras que Ljung-Box rechazó la ausencia conjunta de autocorrelación en los rezagos 1, 4, 13, 26 y 52. No se observó una autocorrelación anual individual fuera del límite aproximado del 95% en el rezago 52. Estos resultados justifican rezagos recientes y la comparación con una referencia estacional, pero no demuestran una estacionalidad anual estable."),
    ]:
        insert_before(baseline_heading, text, style)
    insert_picture_before(baseline_heading, FIGURES / "02_acf_pacf_semanal.png", "Figura AT-2 - Funciones ACF y PACF del bloque de desarrollo")
    insert_picture_before(baseline_heading, FIGURES / "03_perfil_semana_anio.png", "Figura AT-3 - Perfil exploratorio del importe por semana del año")
    for style, text in [
        ("Heading 3", "Intermitencia, variabilidad y valores atípicos"),
        ("Normal", "El bloque de desarrollo contiene 13.5% de semanas con importe cero. El ADI fue 1.16 y el CV² de los importes positivos fue 1.03; de acuerdo con los umbrales establecidos, la serie se clasifica como errática: las compras ocurren con relativa frecuencia, pero sus importes presentan alta variabilidad. Esta caracterización respalda el uso de log1p, modelos hurdle y métodos robustos, y muestra que el problema principal no es una intermitencia extrema sino la dispersión de los montos."),
        ("Normal", "El criterio robusto basado en mediana y MAD identificó siete semanas atípicas dentro del bloque continuo. Estas observaciones se conservaron porque pueden corresponder a compras extraordinarias o eventos comerciales legítimos. Su presencia refuerza la conveniencia de evaluar MAE junto con RMSE y de no seleccionar modelos mediante MAPE."),
    ]:
        insert_before(baseline_heading, text, style)
    insert_picture_before(baseline_heading, FIGURES / "04_atipicos_robustos.png", "Figura AT-4 - Identificación robusta de importes semanales atípicos")
    insert_picture_before(baseline_heading, FIGURES / "05_descomposicion_stl_exploratoria.png", "Figura AT-5 - Descomposición STL exploratoria con periodo de 52 semanas")
    for style, text in [
        ("Heading 3", "Implicaciones para los modelos de Inteligencia Artificial"),
        ("Normal", "Los diagnósticos sustentan un diseño parsimonioso: rezagos recientes, codificación cíclica, transformación log1p, regularización y modelos de dos etapas. La descomposición STL se conserva como visualización, pero sus índices de fuerza no se interpretan porque el bloque de desarrollo contiene sólo 2.13 ciclos anuales. La inteligencia artificial se evalúa así sobre una representación temporal auditada, sin asumir que una mayor complejidad garantizará menor error."),
        ("Normal", "Después de aplicar el rezago máximo de 52 semanas, quedaron 75 observaciones compatibles con el modelado. Se utilizaron 52 semanas por ventana, orígenes previos para ajuste y 16 semanas finales comunes para evaluación. Esta muestra limitada exige interpretar los resultados como evidencia del caso de estudio y no como una generalización a otras microempresas."),
    ]:
        insert_before(baseline_heading, text, style)

    # Actualiza el desarrollo posterior con la nueva cobertura y resultados.
    replace(find_contains(document, "La auditoría examinó cobertura temporal, duplicados"),
        "La auditoría examinó cobertura temporal, duplicados, valores nulos, consistencia de tipos, concentración de ceros y continuidad del registro. El bloque continuo principal quedó delimitado entre 2022-01-03 y 2024-06-03. Las 44 semanas de brecha interna, la cola de 43 semanas y el bloque posterior insuficiente para una ventana de 52 semanas se excluyeron del experimento principal, conservándose como evidencia de calidad y limitación del sistema de captura.")
    replace(find_contains(document, "La ingeniería de características semanal incorpora rezagos"),
        "La ingeniería de características semanal incorpora rezagos de 1, 2, 4, 8 y 52 semanas, medias y desviaciones móviles de 4, 8 y 12 semanas, codificación cíclica y variables exógenas registradas. La selección se realiza dentro de cada ventana de entrenamiento. Para horizontes directos superiores a una semana, las características históricas se congelan en el origen y sólo se admiten exógenas disponibles para la semana objetivo.")
    replace(find_contains(document, "Se comparan referencias empíricas, ETS, ARIMA, Croston-SBA"),
        "Se compararon referencias empíricas, ETS, ARIMA, Croston-SBA, TSB y modelos de aprendizaje automático parsimoniosos. Las variantes log1p y hurdle respondieron a la alta variabilidad y a los ceros observados. RNN y LSTM permanecen como contraste secundario y no como evidencia principal, debido al reducido número de observaciones semanales continuas y al riesgo de sobreajuste.")
    replace(find_contains(document, "La evaluación se ejecuta mediante rolling-window"),
        "La evaluación se ejecutó mediante rolling-window con 52 semanas fijas de entrenamiento, desplazamiento de una semana y 16 orígenes finales comunes dentro del bloque continuo. Las métricas se calcularon por separado para h=1 y h=4. Para el consolidado presupuestal se sumaron los pronósticos directos h=1, h=2, h=3 y h=4; esta suma no se interpreta como un modelo mensual independiente.")
    replace(find_contains(document, "El archivo de salida para el sistema de soporte a la decisión comunica resultados semanales auditables"),
        "El archivo de salida para el sistema de soporte a la decisión comunica resultados semanales auditables: diagnóstico temporal, cobertura utilizable, métricas por horizonte, contrastes H1 y H2, predicciones históricas, consolidaciones de cuatro semanas e inventario de variables exógenas. Su propósito es facilitar la revisión humana del presupuesto y transparentar cuándo la evidencia no respalda una superioridad algorítmica.")
    replace(find_contains(document, "La integración se realiza mediante 04_dss_semanal.json"),
        "La integración se realiza mediante 04_dss_semanal.json, generado a partir de la ejecución con cobertura corregida. El archivo conserva la relación entre cada configuración, las métricas por horizonte, los contrastes estadísticos y las variables exógenas registradas; evita presentar resultados diarios heredados o semanas sin cobertura como evidencia semanal válida.")
    replace(find_contains(document, "Para esta investigación, la salida del DSS se limita al importe semanal de compras"),
        "Para esta investigación, la salida del DSS se limita al importe semanal de compras del bloque con cobertura defendible. Debe mostrar por separado el desempeño en h=1 y h=4, el presupuesto consolidado de cuatro semanas, el diagnóstico de cobertura y las variables exógenas registradas. No debe mezclar resultados diarios heredados ni presentar las brechas como semanas de compras iguales a cero.")
    replace(find_contains(document, "En el horizonte principal h=1, Ridge con variables históricas obtuvo"),
        "En el horizonte principal h=1, la referencia ingenua estacional de 52 semanas obtuvo el menor RMSE descriptivo (695.19) y MAE de 438.19. El modelo de aprendizaje automático más cercano fue Random Forest con transformación log1p y variables históricas, con RMSE de 699.28 y MAE de 478.54. La cercanía descriptiva no constituye evidencia de superioridad y debe leerse junto con los contrastes inferenciales.")
    replace(find_contains(document, "La conclusión del DSS es prudente: ningún contraste frente a la línea base primaria"),
        "Ninguna configuración estadística o de aprendizaje automático superó significativamente la línea base primaria de último valor después del ajuste Holm, tanto en h=1 como en h=4; por ello H1 no quedó respaldada. H2 mostró mejoras dependientes de la configuración —por ejemplo, Ridge enriquecido en h=1 obtuvo una diferencia favorable sin ajuste múltiple—, pero no autoriza afirmar que las variables exógenas mejoren universalmente todos los modelos.")
    replace(find_contains(document, "El DSS debe apoyar la revisión del presupuesto y mostrar la incertidumbre asociada"),
        "El DSS debe apoyar la revisión del presupuesto y mostrar la incertidumbre asociada a cada horizonte. En h=1 se comunica la referencia estacional como menor RMSE descriptivo; en h=4, HistGradientBoosting tipo hurdle con variables exógenas obtuvo el menor RMSE descriptivo (399.56) y MAE de 318.52. Ninguno de estos resultados autoriza automatizar compras ni afirmar superioridad estadísticamente demostrada frente a la referencia primaria.")
    replace(find_contains(document, "La cobertura se presenta como una condición de uso: después de depurar la cola sin observación"),
        "La cobertura se presenta como condición de uso. El bloque principal contiene 127 semanas continuas; después del rezago máximo quedaron 75 observaciones para modelado y 16 orígenes finales para evaluación. La brecha interna, la cola sin cobertura y el segmento posterior insuficiente se mantienen fuera del ajuste y de los contrastes. Antes de utilizar el sistema operativamente deben reanudarse y verificarse los registros de compras.")

    # Referencia nueva requerida por ADI/CV² y SBA.
    reference_anchor = find_contains(document, "Tashman, L. J. (2000)")
    insert_before(reference_anchor,
        "Syntetos, A. A., & Boylan, J. E. (2005). The accuracy of intermittent demand estimates. International Journal of Forecasting, 21(2), 303–314. https://doi.org/10.1016/j.ijforecast.2004.10.001",
        reference_anchor.style.name)

    enable_field_updates(document)
    if document.paragraphs[3].text.strip() != EXPECTED_TITLE:
        raise ValueError("El título fue modificado accidentalmente.")
    document.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
