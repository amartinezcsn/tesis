from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev38_(ZUJ)_22ago2026_Tabla1ColumnasCompletas.docx")
OUTPUT = Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev39_(ZUJ)_22ago2026_JustificacionAlineada.docx")


def normalized(value):
    return " ".join(value.split())


def find_prefix(document, prefix):
    prefix = normalized(prefix)
    matches = [p for p in document.paragraphs if normalized(p.text).startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph beginning with {prefix!r}; found {len(matches)}")
    return matches[0]


def replace(paragraph, text):
    props = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if props is not None:
        run._element.insert(0, props)


def insert_after(document, anchor, text):
    paragraph = document.add_paragraph(style=anchor.style)
    if anchor._p.pPr is not None:
        paragraph._p.get_or_add_pPr().getparent().replace(paragraph._p.get_or_add_pPr(), deepcopy(anchor._p.pPr))
    paragraph.add_run(text)
    anchor._p.addnext(paragraph._p)
    return paragraph


doc = Document(SOURCE)

replace(
    find_prefix(doc, "La relevancia social de esta investigación radica"),
    "La relevancia social de esta investigación radica en aportar evidencia aplicada sobre el uso de aprendizaje automático en una microempresa con recursos de datos limitados. Las microempresas enfrentan restricciones financieras, tecnológicas y de capital humano que dificultan la adopción de herramientas analíticas; por ello, evaluar modelos predictivos con registros internos de ventas y compras contribuye a identificar alternativas técnicamente accesibles para apoyar decisiones basadas en datos (Alekseeva et al., 2021; Poveda-Valverde & Fierro Barragán, 2026).",
)
replace(
    find_prefix(doc, "Al mostrar la viabilidad de aplicar modelos de aprendizaje automático"),
    "En el caso de Cup&Cake, la investigación permite valorar si el aprendizaje automático mejora la anticipación del presupuesto de abastecimiento frente a referencias empíricas y estadísticas. Aunque el estudio no mide directamente merma, inventarios ni rentabilidad, una predicción más precisa del importe de compras puede constituir un insumo para decisiones más informadas en negocios con productos perecederos.",
)
replace(
    find_prefix(doc, "La presente investigación resulta conveniente"),
    "La investigación es conveniente porque responde a una necesidad operativa concreta de Cup&Cake: anticipar el importe diario de compras a partir de sus registros históricos de ventas y compras. La disponibilidad de estas fuentes permite construir y evaluar un sistema predictivo sin requerir infraestructura corporativa ni datos externos indispensables.",
)
replace(
    find_prefix(doc, "En este escenario, la toma de decisiones basada únicamente"),
    "Su conveniencia metodológica consiste en comparar modelos de aprendizaje automático con una línea base empírica reproducible y con modelos estadísticos de referencia. Este contraste permite identificar si la mayor complejidad algorítmica aporta una mejora verificable de precisión o si un método más simple resulta suficiente para el objetivo evaluado (Makridakis et al., 2018; Spiliotis et al., 2022).",
)
replace(
    find_prefix(doc, "Desde el punto de vista teórico, la investigación contribuye"),
    "Desde el punto de vista teórico, la investigación contribuye al estudio del pronóstico en contextos de Small Data e intermitencia. La literatura muestra que el desempeño de los modelos de aprendizaje automático depende de la calidad y cantidad de datos, la granularidad de la serie, las variables disponibles y el protocolo de evaluación; por tanto, su superioridad no debe suponerse sin comparación empírica (Hewamalage et al., 2021; Giannopoulos et al., 2025).",
)
replace(
    find_prefix(doc, "El caso propuesto permite discutir hasta qué punto"),
    "El estudio aporta evidencia comparativa en una microempresa de repostería creativa, al contrastar una línea base empírica, modelos estadísticos y modelos de aprendizaje automático mediante validación temporal Rolling-Origin. Así, delimita en qué condiciones el aprendizaje automático puede aportar valor para anticipar el presupuesto de abastecimiento.",
)
replace(
    find_prefix(doc, "En términos prácticos, ofrece un procedimiento replicable"),
    "En términos prácticos, la investigación ofrece un procedimiento reproducible para integrar registros históricos de ventas y compras, generar variables temporales y evaluar modelos predictivos mediante métricas de error fuera de muestra. Los resultados se comunican en un tablero de inteligencia de negocios que permite consultar el importe y la frecuencia esperada de compras, así como el desempeño de cada alternativa frente a la línea base.",
)
last = find_prefix(doc, "Adicionalmente, la comparación entre el método empírico actual")
replace(
    last,
    "El tablero funciona como apoyo para revisar y reservar el presupuesto de abastecimiento; no genera órdenes automáticas de compra ni calcula cantidades físicas por insumo. Estas funciones requerirían información adicional sobre recetas, inventario, merma, costos de faltante, tiempos de entrega y unidades de medida homologadas.",
)
insert_after(
    doc,
    last,
    "En conjunto, la investigación se justifica por su aporte aplicado y metodológico: evalúa de manera rigurosa si los modelos de aprendizaje automático mejoran la predicción del importe diario de compras en una microempresa, frente a referencias empíricas y estadísticas. Su contribución se limita al apoyo del presupuesto de abastecimiento y evita atribuir efectos no medidos sobre inventarios, merma, rentabilidad o expansión comercial.",
)

doc.core_properties.title = "Predicción del presupuesto de abastecimiento mediante aprendizaje automático"
doc.save(OUTPUT)
print(OUTPUT)
