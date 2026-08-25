from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

p=Path(r"C:\Python\tesis\documentacion\TESIS_AGO2026_Rev33_(ZUJ)_21ago2026.docx")
d=Document(p)

# La codificación cíclica ya estaba expresada en una ecuación nativa del marco
# teórico; se conserva una sola aparición para evitar duplicidad conceptual.
for table in list(d.tables):
    if 'xₛₑₙ' in ''.join(c.text for row in table.rows for c in row.cells):
        nxt=table._tbl.getnext()
        table._tbl.getparent().remove(table._tbl)
        if nxt is not None:
            txt=''.join(nxt.itertext()).strip()
            if txt.startswith('Ecuación - Codificación cíclica'):
                nxt.getparent().remove(nxt)
        break

# Elimina las fórmulas métricas antiguas que habían quedado duplicadas en el marco teórico.
for idx,par in list(enumerate(d.paragraphs)):
    if idx < 500 and par.style.name=='Caption' and par.text.strip().startswith('Ecuación'):
        prev=par._p.getprevious()
        if prev is not None and prev.findall('.//'+qn('m:oMath')):
            prev.getparent().remove(prev)
        par._element.getparent().remove(par._element)

# Las leyendas descriptivas de las nuevas ecuaciones no participan en la secuencia;
# la numeración automática permanece en la tercera celda de cada tabla.
for par in d.paragraphs:
    if par.style.name=='Caption' and par.text.strip().startswith('Ecuación -'):
        par.style=d.styles['Normal']; par.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in par.runs: r.italic=True

# Todos los capítulos principales comienzan en página nueva.
for par in d.paragraphs:
    if par.style.name=='Heading 1': par.paragraph_format.page_break_before=True
    if par.style.name=='Heading 1' and 'HYPERLINK' in par.text and 'CONCLUSIONES' in par.text:
        par.clear(); par.add_run('CONCLUSIONES')
    if par.style.name=='Heading 1':
        codes=[''.join(x.itertext()) for x in par._p.findall('.//'+qn('w:instrText'))]
        if any('HYPERLINK' in x for x in codes):
            visible=''.join(x.text or '' for x in par._p.findall('.//'+qn('w:t'))).strip()
            par.clear(); par.add_run(visible or 'DISCUSION')

# Retira un encabezado vacío heredado del documento fuente.
for par in list(d.paragraphs):
    if par.style.name.startswith('Heading') and not par.text.strip():
        par._element.getparent().remove(par._element)

# Los campos ya fueron materializados en Word durante el control final. Se evita
# una actualización automática adicional al abrir, pues el documento fuente
# contiene marcadores heredados que Word intenta reparar dentro de los títulos.
upd=d.settings._element.find(qn('w:updateFields'))
if upd is None:
    upd=OxmlElement('w:updateFields'); d.settings._element.append(upd)
upd.set(qn('w:val'),'false')

d.save(p)
print(p)
