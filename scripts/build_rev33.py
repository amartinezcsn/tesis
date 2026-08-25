from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(r"C:\Python\tesis")
DOC=ROOT/r"documentacion\TESIS_AGO2026_Rev33_(ZUJ)_21ago2026.docx"
IMG=ROOT/"imagenes"
d=Document(DOC)

def find(text):
    for p in d.paragraphs:
        if p.text.strip()==text.strip(): return p
    raise ValueError(text)

def find_contains(text):
    for p in d.paragraphs:
        if text.lower() in p.text.lower(): return p
    raise ValueError(text)

def new_p_after(anchor, text="", style=None, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p=d.add_paragraph(style=style)
    r=p.add_run(text); r.bold=bold; r.italic=italic
    p.alignment=align; p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
    anchor._p.addnext(p._p)
    return p

def add_series(anchor, items):
    cur=anchor
    for item in items:
        if isinstance(item,tuple):
            text,style=item
        else: text,style=item,None
        cur=new_p_after(cur,text,style)
    return cur

def set_text(p,text):
    p.clear(); p.add_run(text); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def delete_between(start, end):
    el=start._p.getnext(); stop=end._p
    while el is not None and el is not stop:
        nxt=el.getnext(); el.getparent().remove(el); el=nxt

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def add_table(anchor, title, headers, rows, widths=None):
    cap=new_p_after(anchor, title, 'Caption', align=WD_ALIGN_PARAGRAPH.CENTER)
    t=d.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c,'D9EAF7'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    cap._p.addnext(t._tbl)
    return t

def add_picture(anchor, filename, caption, width=6.3):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(IMG/filename), width=Inches(width)); anchor._p.addnext(p._p)
    cap=new_p_after(p,caption,'Caption',align=WD_ALIGN_PARAGRAPH.CENTER)
    return cap

def field_run(p,instr,result='1'):
    r=OxmlElement('w:r'); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); r.append(b); p._p.append(r)
    r=OxmlElement('w:r'); it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr; r.append(it); p._p.append(r)
    r=OxmlElement('w:r'); s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); r.append(s); p._p.append(r)
    p.add_run(result)
    r=OxmlElement('w:r'); e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r.append(e); p._p.append(r)

def math_element(text):
    omp=OxmlElement('m:oMathPara'); om=OxmlElement('m:oMath'); omp.append(om)
    mr=OxmlElement('m:r'); mt=OxmlElement('m:t'); mt.text=text; mr.append(mt); om.append(mr)
    return omp

eq_template=deepcopy(d.tables[1]._tbl)
def add_equation(anchor, formula, label):
    tbl=deepcopy(eq_template); cells=tbl.findall('.//w:tr/w:tc',{'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})[:3]
    for c in cells:
        for ch in list(c):
            if ch.tag!=qn('w:tcPr'): c.remove(ch)
        p=OxmlElement('w:p'); c.append(p)
    cells[1].find(qn('w:p')).append(math_element(formula))
    rp=cells[2].find(qn('w:p')); pp=type('P',(),{'_p':rp,'add_run':lambda self,x: __import__('docx').text.paragraph.Paragraph(rp,rp.getparent()).add_run(x)})()
    para=__import__('docx').text.paragraph.Paragraph(rp,rp.getparent()); para.alignment=WD_ALIGN_PARAGRAPH.RIGHT; para.add_run('('); field_run(para,' SEQ Ecuacion \\* ARABIC ','1'); para.add_run(')')
    anchor._p.addnext(tbl)
    cap=d.add_paragraph(style='Caption'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.add_run(label); tbl.addnext(cap._p)
    return cap

# Correcciones editoriales y metodológicas puntuales
for p in d.paragraphs:
    if '28 de agosto de 2025' in p.text: set_text(p,p.text.replace('28 de agosto de 2025','21 de agosto de 2026'))
    if 'ventas e compras' in p.text: set_text(p,p.text.replace('ventas e compras','ventas y compras'))
    if 'Power BI' in p.text or 'en Un tablero' in p.text: set_text(p,p.text.replace('en Un tablero','en un tablero').replace('Power BI','un tablero web de inteligencia de negocios'))
    if 'TURNITY' in p.text: set_text(p,p.text.replace('TURNITY','Cup&Cake'))

# Resumen actualizado y Abstract
res=find('RESUMEN'); intro=find('INTRODUCCIÓN'); delete_between(res,intro)
cur=add_series(res,[
"La investigación desarrolló y evaluó un sistema de pronóstico para apoyar la planeación financiera y el abastecimiento de Cup&Cake, microempresa de repostería ubicada en Tizayuca, Hidalgo. Se integraron registros diarios de ventas y compras con variables calendáricas, económicas y climáticas; posteriormente se construyeron rezagos, ventanas móviles y características cíclicas bajo controles de fuga de información. La evaluación se realizó mediante validación temporal Rolling-Origin y comparación contra una línea base empírica.",
"El dataset maestro quedó conformado por 1,612 observaciones y 65 variables, mientras que el conjunto de modelado reunió 1,584 días y 276 columnas. Los modelos ganadores fueron ARIMA para los objetivos de ventas y Random Forest para los objetivos de compras. La reducción relativa del RMSE fue de 3.2 % y 5.1 % en ventas, y de 28.3 % y 22.8 % en compras; por ello, la hipótesis se aceptó parcialmente. Los resultados se integraron en un sistema de soporte a la decisión con indicadores, filtros y salvaguardas de interpretación humana.",
"Palabras clave: pronóstico de demanda, microempresa, series temporales, aprendizaje automático, Rolling-Origin, inteligencia de negocios.",
('ABSTRACT','Heading 2'),
"This research developed and evaluated a forecasting system to support financial planning and procurement at Cup&Cake, a small bakery located in Tizayuca, Hidalgo. Daily sales and purchasing records were integrated with calendar, economic and weather variables. Lagged variables, moving windows and cyclical encodings were generated under information-leakage controls. Models were assessed through Rolling-Origin temporal validation against an empirical baseline. ARIMA achieved the best performance for sales targets, whereas Random Forest was selected for purchasing targets. Relative RMSE improvements were 3.2% and 5.1% for sales, and 28.3% and 22.8% for purchases; therefore, the hypothesis was partially supported. Results were incorporated into a business-intelligence decision-support dashboard.",
"Keywords: demand forecasting, microenterprise, time series, machine learning, Rolling-Origin, business intelligence."
])

# Metodología: diseño coherente
h=find_contains('Diseño de la Investigación: Experimento'); h.text='Diseño evaluativo y comparación contra la línea base'
next_h=find('Según el tipo de inferencia'); delete_between(h,next_h)
add_series(h,[
"Se adoptó un diseño cuantitativo, aplicado, no experimental y evaluativo, basado en un estudio de caso longitudinal-retrospectivo. No se manipuló deliberadamente la operación de la empresa ni se asignaron grupos aleatorios; se analizaron registros históricos ordenados cronológicamente.",
"La preprueba se operacionalizó como el error obtenido por el método empírico de referencia —último valor disponible y promedio móvil de siete días—. La posprueba correspondió al error de los modelos predictivos sobre las mismas ventanas futuras. En consecuencia, el término control se refiere a una línea base analítica y no a un grupo de control experimental.",
"La comparación utilizó tres orígenes temporales finales. En cada iteración, el entrenamiento incluyó exclusivamente observaciones anteriores al periodo de prueba, preservando la secuencia causal y evitando fuga de información."
])
tp=find_contains('Según el periodo temporal'); tp.text='Temporalidad de la investigación'
proc=find('Procedimiento Metodológico'); delete_between(tp,proc)
add_series(tp,["La investigación fue longitudinal y retrospectiva, porque examinó la evolución diaria de ventas y compras entre 2022 y 2026 mediante observaciones repetidas del mismo negocio. Esta temporalidad permitió estudiar estacionalidad, cambios de nivel, periodos sin actividad y estabilidad del error predictivo."])

# Aplanamiento de jerarquía excesiva
for p in d.paragraphs:
    if p.style.name in ('Heading 5','Heading 6'):
        p.style=d.styles['Heading 4'] if not p.text.startswith('Etapa ') else d.styles['Normal'];
        if p.text.startswith('Etapa '): p.runs[0].bold=True if p.runs else None

# Sustituir bloque de recomendación técnica por síntesis científica
rec=find('Recomendación para la tesis'); diag=find_contains('Diagnóstico de la Línea Base'); delete_between(rec,diag); rec._element.getparent().remove(rec._element)

# Secciones de desarrollo previamente vacías
aud=find_contains('Auditoría, Depuración'); fe=find_contains('Ingeniería de Características (Feature Engineering) y'); mod=find_contains('Desarrollo, Configuración'); eva=find_contains('Evaluación Comparativa'); dss=find_contains('Arquitectura e Integración')
add_series(diag,["La línea base representó la práctica empírica reproducible con la que se contrastaron los modelos. Para cada objetivo se evaluaron la persistencia del último valor y el promedio móvil de siete días; se conservó como referencia el método con menor RMSE en cada ventana temporal. Este planteamiento permitió cuantificar la ganancia analítica sin atribuir causalidad experimental."])
add_table(diag,'Tabla 7 - Operacionalización de la línea base y la posprueba',['Componente','Definición operacional','Indicador'],[
['Preprueba','Método empírico: último valor o promedio de 7 días','RMSE y MAE de referencia'],['Posprueba','Mejor modelo evaluado en los mismos orígenes','RMSE y MAE del modelo'],['Contraste','Reducción relativa del RMSE','Umbral de H1: 20 %']],[1.2,3.5,1.6])
add_series(aud,["La auditoría examinó cobertura temporal, duplicados, valores nulos, consistencia de tipos, concentración de ceros y continuidad diaria. Después de homologar fechas, categorías y unidades monetarias reales, las fuentes se integraron por fecha en un dataset maestro trazable. El pipeline completo finalizó sin errores y produjo artefactos de control para cada etapa."])
add_table(aud,'Tabla 8 - Trazabilidad de los conjuntos analíticos',['Conjunto','Observaciones','Variables','Cobertura o criterio'],[
['Dataset maestro','1,612','65','2022-01-01 a 2026-05-31; sin fechas faltantes'],['Dataset modelado','1,584','276','2022-01-29 a 2026-05-31'],['Dataset reducido','1,584','85','80 predictores, fecha y cuatro objetivos'],['Dataset PCA','1,584','135','130 componentes, fecha y cuatro objetivos']],[1.5,1.1,1.0,3.0])
add_series(fe,["La ingeniería de características transformó las series originales en una representación supervisada de 271 predictores. Se incorporaron calendario, codificación cíclica, eventos exógenos, rezagos, ventanas móviles, composición por categorías, estabilidad y recencia. Todas las características históricas se desplazaron temporalmente y se eliminaron los primeros 28 días para asegurar ventanas completas."])
pic=add_picture(fe,'24_expansion_dimensional_feature_engineering.png','Figura 24 - Expansión dimensional producida por la ingeniería de características.')
add_table(pic,'Tabla 9 - Dimensiones principales de las variables modeladas',['Dimensión','Número de variables','Finalidad analítica'],[
['Rezagos','126','Memoria temporal de corto y mediano plazo'],['Ventanas móviles','105','Nivel, dispersión y acumulación reciente'],['Calendario y estacionalidad','19','Ciclos semanales, mensuales y anuales'],['Otras variables','21','Exógenas, relaciones, estabilidad y objetivos']],[2.0,1.3,3.3])
add_series(mod,["Se compararon enfoques empíricos, ARIMA/SARIMA, regresión lineal, árboles, Random Forest, RNN y LSTM. Cada familia se entrenó sobre representaciones completa, reducida y PCA cuando resultó procedente. Los hiperparámetros se seleccionaron dentro del conjunto de entrenamiento de cada origen temporal; la comparación final privilegió RMSE, MAE, estabilidad entre ventanas e interpretabilidad."])
add_table(mod,'Tabla 10 - Familias de modelos y configuración evaluada',['Familia','Configuración general','Representaciones'],[
['Empírica','Persistencia y promedio móvil de 7 días','Serie original'],['ARIMA/SARIMA','Órdenes seleccionados por objetivo','Completa y univariada'],['Regresión y árboles','Lineal, árbol de decisión y Random Forest','Completa, reducida y PCA'],['Redes recurrentes','RNN simple y LSTM con ventanas temporales','Reducida y PCA']],[1.5,3.2,2.0])
pic=add_picture(mod,'03_arquitectura_rnn_lstm.png','Figura 25 - Arquitecturas recurrentes consideradas para el modelado temporal.')
add_series(eva,["La evaluación se ejecutó mediante Rolling-Origin. Los conjuntos de entrenamiento crecieron de forma acumulativa y cada prueba se situó después del último dato empleado para ajustar el modelo. La selección global se basó en el promedio del error y su variabilidad en los tres orígenes finales."])
pic=add_picture(eva,'01_esquema_validacion_rolling_origin.png','Figura 26 - Esquema de validación temporal Rolling-Origin.')
add_equation(pic,'RMSĒ = (1/K) Σₖ RMSEₖ ;  Entrenamientoₖ = {t ≤ τₖ},  Pruebaₖ = {τₖ < t ≤ τₖ+h}','Ecuación - Formalización de la validación Rolling-Origin.')

# Figuras científicas en integración y DSS
add_picture(find('El proceso de construcción del dataset maestro'),'88_pipeline_trazabilidad_tesis.png','Figura 27 - Trazabilidad del pipeline analítico de la investigación.')
add_picture(dss,'90_arquitectura_funcional_dss.png','Figura 28 - Arquitectura funcional del sistema de soporte a la decisión.')
alc=find_contains('Alcance actual de la integración'); add_picture(alc,'87_captura_tablero_dss.png','Figura 29 - Interfaz funcional del tablero DSS implementado para Cup&Cake.')
sim=find_contains('Simulación de Escenarios'); sim.text='Alcance de la simulación financiera y de abastecimiento'
results=find('RESULTADOS'); delete_between(sim,results)
add_series(sim,["El DSS permite explorar el comportamiento de los cuatro objetivos, comparar el desempeño contra la línea base y revisar estabilidad por ventana. Sin embargo, no se ejecutó una optimización prescriptiva de inventario porque los registros disponibles no contienen de manera consistente costos de faltante, merma, inventario inicial ni tiempos de entrega. Por rigor científico, los escenarios se interpretaron como apoyo diagnóstico y no como órdenes automáticas de compra."])
add_picture(sim,'91_flujo_decision_operativa_dss.png','Figura 30 - Flujo de decisión operativa asistido por el DSS.')

# Ecuaciones adicionales en lugares metodológicamente pertinentes
np=find_contains('Problema del Vendedor de Periódicos'); add_equation(np,'F(Q*) = Cᵤ / (Cᵤ + Cₒ)','Ecuación - Fracción crítica del modelo Newsvendor.')
cy=find_contains('Codificación cíclica de las variables temporales'); add_equation(cy,'xₛₑₙ = sin(2πx/P) ;  x꜀ₒₛ = cos(2πx/P)','Ecuación - Codificación cíclica de una variable temporal.')
hist=find_contains('Características históricas y composición'); add_equation(hist,'lagₖ(yₜ)=yₜ₋ₖ ;  MA_w(yₜ)=(1/w)Σᵢ₌₁ʷ yₜ₋ᵢ','Ecuación - Rezagos y media móvil sin información futura.')
mult=find_contains('Etapa 2. Diagnóstico de multicolinealidad'); add_equation(mult,'VIFⱼ = 1 / (1 − Rⱼ²)','Ecuación - Factor de inflación de la varianza.')
pca=find_contains('Etapa 4. Reducción'); add_equation(pca,'Z = XW ;  Varianza acumulada(m)=Σⱼ₌₁ᵐ λⱼ / Σⱼ₌₁ᵖ λⱼ','Ecuación - Transformación PCA y varianza explicada acumulada.')
inpc=find_contains('Variables Exógenas Inflación'); add_equation(inpc,'Importe realₜ = Importe nominalₜ × (INPC_base / INPCₜ)','Ecuación - Deflactación de importes monetarios.')

# Resultados completos
discussion=find('DISCUSION'); delete_between(results,discussion)
cur=add_series(results,[
('Calidad, cobertura e integración de datos','Heading 2'),
"El pipeline se completó sin errores. El dataset maestro reunió 1,612 días consecutivos entre el 1 de enero de 2022 y el 31 de mayo de 2026, con 65 columnas, sin fechas faltantes ni valores nulos. Después de exigir 28 días de historia para rezagos y ventanas, el dataset de modelado conservó 1,584 observaciones. La cobertura de ventas llegó a mayo de 2026; la de compras terminó el 30 de julio de 2025, por lo que 305 días posteriores se excluyeron de la evaluación de esos objetivos.",
('Reducción dimensional','Heading 2'),
"El perfil inicial identificó 271 predictores. La selección supervisada conservó 80 variables y produjo un dataset reducido de 85 columnas al añadir fecha y cuatro objetivos. Como representación alternativa, PCA necesitó 130 componentes para explicar al menos 95 % de la varianza, lo que confirma que la estructura informativa se encontraba distribuida en numerosas dimensiones.",
('Rendimiento comparativo de los modelos','Heading 2'),
"ARIMA fue la familia ganadora para ambos objetivos de ventas, mientras que Random Forest sobre el dataset reducido obtuvo el mejor desempeño en compras. Los resultados se resumen en la tabla siguiente; la MAPE se reporta únicamente como referencia debido a la elevada frecuencia de valores cero."])
add_table(cur,'Tabla 11 - Modelos ganadores y desempeño temporal',['Objetivo','Modelo y representación','MAE','RMSE','Mejora RMSE'],[
['Importe de ventas','ARIMA(1,1,2)','101.951','161.149','3.2 %'],['Registros de ventas','ARIMA(1,1,1)','0.362','0.525','5.1 %'],['Importe de compras','Random Forest, reducido','116.961','213.014','28.3 %'],['Registros de compras','Random Forest, reducido','1.727','2.593','22.8 %']],[1.7,2.2,1.0,1.0,1.1])
add_equation(find_contains('Rendimiento comparativo de los modelos'),'MAE = (1/n) Σᵢ₌₁ⁿ |yᵢ − ŷᵢ|','Ecuación - Error absoluto medio.')
add_equation(find_contains('Rendimiento comparativo de los modelos'),'RMSE = √[(1/n) Σᵢ₌₁ⁿ (yᵢ − ŷᵢ)²]','Ecuación - Raíz del error cuadrático medio.')
add_equation(find_contains('Rendimiento comparativo de los modelos'),'MAPE = (100/|I|) Σᵢ∈I |(yᵢ − ŷᵢ)/yᵢ|,  I={i:yᵢ≠0}','Ecuación - Error porcentual absoluto medio sobre observaciones no nulas.')
add_equation(find_contains('Rendimiento comparativo de los modelos'),'WAPE = 100 × Σᵢ |yᵢ − ŷᵢ| / Σᵢ |yᵢ|','Ecuación - Error porcentual absoluto ponderado como diagnóstico complementario.')
add_equation(find_contains('Rendimiento comparativo de los modelos'),'Mejora_RMSE = 100 × (RMSE_base − RMSE_modelo) / RMSE_base','Ecuación - Reducción relativa del RMSE frente a la línea base.')
pic=add_picture(find_contains('Rendimiento comparativo de los modelos'),'05_mejora_frente_linea_base_empirica.png','Figura 31 - Mejora relativa del RMSE frente a la línea base empírica.')
pic=add_picture(pic,'04_mejor_modelo_por_dataset_objetivo.png','Figura 32 - Mejor modelo por objetivo y representación de datos.')
cur=add_series(pic,[('Resultados de redes neuronales recurrentes','Heading 2'),"Las arquitecturas RNN y LSTM no superaron a los modelos ganadores. El mejor RMSE recurrente fue 163.777 para importe de ventas, 0.545 para registros de ventas, 239.714 para importe de compras y 2.748 para registros de compras. El resultado respalda el uso de modelos más parsimoniosos ante el tamaño muestral disponible y la demanda intermitente.",('Contraste de hipótesis','Heading 2'),"La hipótesis H1 estableció una reducción mínima de 20 % en el error respecto de la línea base. El umbral se alcanzó en importe y registros de compras, pero no en los dos objetivos de ventas. En consecuencia, H1 se aceptó parcialmente: la evidencia apoya una mejora sustantiva para el abastecimiento, no una mejora generalizada para todos los procesos.",('Resultados del tablero DSS','Heading 2'),"El tablero integró los modelos ganadores, métricas globales, desempeño por ventana, cobertura y contraste de hipótesis. Los filtros permiten seleccionar objetivo y representación; las salvaguardas advierten cuando la cobertura está desactualizada o la mejora no alcanza el umbral. La interfaz funciona como capa de comunicación analítica y mantiene la decisión final bajo responsabilidad del usuario."])

# Discusión
concl=find('CONCLUSIONES'); delete_between(discussion,concl)
add_series(discussion,[
('Interpretación de los hallazgos','Heading 2'),
"La superioridad de ARIMA en ventas indica que la dinámica agregada se explica mejor mediante dependencia temporal y diferenciación que mediante una estructura de alta complejidad. En compras, Random Forest capturó relaciones no lineales entre rezagos, ventanas y composición categórica, con mejoras superiores al umbral de la hipótesis.",
"El resultado de RNN y LSTM muestra que una arquitectura más compleja no implica automáticamente una mejor predicción. Con 1,584 observaciones, alta dimensionalidad y numerosos días con valor cero, el incremento de parámetros eleva el riesgo de varianza y limita la generalización.",
('Demanda intermitente y selección de métricas','Heading 2'),
"La proporción de días con valor cero fue aproximadamente 75.9 % en ventas y 82.0 % en compras. Esta condición vuelve inestable la MAPE y explica sus valores elevados; por ello, las conclusiones se fundamentaron principalmente en RMSE y MAE. WAPE se incorporó como métrica complementaria propuesta para futuras corridas, sin atribuirle resultados que no fueron persistidos por el pipeline.",
('Implicaciones operativas y financieras','Heading 2'),
"La evidencia permite priorizar el uso predictivo en compras, donde la mejora fue material y consistente con la necesidad de reducir faltantes y sobreabastecimiento. Para ventas, el DSS debe emplearse como referencia de tendencia y alerta, acompañado del juicio del responsable del negocio, ya que la ganancia frente a la práctica empírica fue modesta.",
('Validez y limitaciones','Heading 2'),
"La validez interna se fortaleció mediante separación temporal, desplazamiento de características y comparación sobre ventanas equivalentes. La validez externa es limitada porque se estudió una sola microempresa. Además, la menor cobertura de compras, la concentración de ceros y la ausencia de inventarios, mermas, tiempos de entrega y costos de faltante restringen la simulación prescriptiva."
])

# Conclusiones, recomendaciones y trabajo futuro
gloss=find('GLOSARIO'); refs=find('REFERENCIAS'); delete_between(concl,gloss)
add_series(concl,[
"Se construyó una solución analítica reproducible que integra datos transaccionales y exógenos, transforma las series mediante ingeniería de características, compara modelos con validación temporal y comunica los resultados en un DSS. El proceso convirtió registros dispersos en evidencia verificable para la planeación financiera y de abastecimiento.",
"La hipótesis se aceptó parcialmente. Random Forest redujo el RMSE de compras en 28.3 % para importes y 22.8 % para registros; ARIMA mejoró ventas en 3.2 % y 5.1 %, por debajo del umbral de 20 %. Por tanto, el principal valor operativo se concentra en el apoyo al abastecimiento.",
"El estudio también demuestra que, en un contexto de Small Data y demanda intermitente, la parsimonia, la calidad de las variables y la validación temporal pueden ser más relevantes que la complejidad algorítmica.",
('Recomendaciones','Heading 2'),
"Se recomienda actualizar mensualmente las fuentes, vigilar la cobertura de compras, conservar la línea base como control permanente y reentrenar cuando el RMSE por ventana muestre deterioro sostenido. Las decisiones de compra deben incorporar revisión humana y límites financieros.",
('Trabajo futuro','Heading 2'),
"Conviene registrar inventario inicial y final, merma, costo de faltante, tiempo de entrega y promociones; evaluar modelos específicos para demanda intermitente; calcular WAPE de forma persistente; y efectuar una validación prospectiva antes de automatizar recomendaciones de abastecimiento."
])
gloss.style=d.styles['Heading 1']; refs.style=d.styles['Heading 1']

# Ajustes de campos y metadatos
settings=d.settings._element; upd=settings.find(qn('w:updateFields'))
if upd is None: upd=OxmlElement('w:updateFields'); settings.append(upd)
upd.set(qn('w:val'),'true')
d.core_properties.title='Sistema predictivo y DSS para la planeación financiera y de abastecimiento de Cup&Cake'
d.core_properties.subject='Tesis de maestría — versión revisada Rev33'
d.core_properties.comments='Revisión estructural, metodológica y visual aprobada; ecuaciones nativas numeradas automáticamente.'
d.save(DOC)
print(DOC)
